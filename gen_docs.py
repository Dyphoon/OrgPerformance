#!/usr/bin/env python3
# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

def set_cell_bg(cell, color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)

def add_heading(doc, text, level=1, color=None):
    p = doc.add_heading(text, level=level)
    if color:
        for run in p.runs:
            run.font.color.rgb = RGBColor(*color)
    return p

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # header
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        set_cell_bg(cell, '2E4057')
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(10)

    # data
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        bg = 'F8F9FA' if ri % 2 == 0 else 'FFFFFF'
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = str(val)
            set_cell_bg(cell, bg)
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.font.size = Pt(9)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    return table

doc = Document()

# 页边距
from docx.oxml import OxmlElement
section = doc.sections[0]
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)

# 标题
title = doc.add_heading('机构绩效管理系统\n产品需求文档', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph(f'版本：V1.0.0\n日期：{datetime.date.today().strftime("%Y-%m-%d")}\n状态：初稿').alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

# 目录（简化）
add_heading(doc, '目录', 1)
doc.add_paragraph('1. 项目概述\n2. 用户与角色\n3. 功能需求\n4. 考核体系管理\n5. 月度监测流程\n6. 数据收集\n7. 结果展示\n8. 通知模块\n9. 文件存储\n10. 非功能需求')

doc.add_page_break()

add_heading(doc, '1. 项目概述', 1)
doc.add_heading('1.1 项目背景', 2)
doc.add_paragraph(
    '随着银行规模的不断扩大，下属分行数量增多，绩效考核的管理复杂度显著提升。'
    '传统人工汇总Excel、手动计算的方式效率低、易出错，且数据分散难以形成统一分析。'
    '机构绩效管理系统旨在通过标准化的Excel模版驱动，实现绩效考核全流程的线上化管理。'
)

doc.add_heading('1.2 项目目标', 2)
doc.add_paragraph(
    '建立一套覆盖考核体系创建、月度数据收集、绩效结果计算与可视化展示的完整系统，'
    '实现多分行绩效考核的规范化、自动化、可视化管理，提升绩效管理效率与数据准确性。'
)

doc.add_heading('1.3 核心价值', 2)
items = [
    '标准统一：通过Excel模版固化考核体系，确保各分行口径一致',
    '流程可控：月度监测全流程线上化，状态透明，节点可追溯',
    '数据准确：指标由Excel公式计算，系统只读取结果，消除人工误差',
    '高效协同：支持批量分发任务，收数人独立填写，管理员统一管理',
    '数据驱动：绩效结果可视化，支持分组对比与历史追溯',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()
add_heading(doc, '2. 用户与角色', 1)

doc.add_heading('2.1 角色定义', 2)
add_table(doc,
    ['角色', '说明', '主要操作'],
    [
        ['绩效管理员', '银行总行绩效管理人员，负责考核体系的创建与维护，发起和管理月度监测流程', '创建体系、上传模版、发起监测、审批数据、发布结果'],
        ['收数人', '各分行具体负责数据填报的人员', '接收任务、填写数据、提交数据'],
        ['机构负责人', '各分行领导，有权限查看本机构及所有机构的绩效报表', '确认数据、查看报表、下载文件'],
    ],
    col_widths=[2.5, 5, 5]
)

doc.add_heading('2.2 菜单权限', 2)
add_table(doc,
    ['菜单', '绩效管理员', '收数人', '机构负责人'],
    [
        ['考核体系管理', '✓', '-', '-'],
        ['月度考核监测', '✓', '✓（仅数据填写）', '-'],
        ['绩效结果展示', '✓', '-', '✓（发布后）'],
    ],
    col_widths=[3, 2.5, 2.5, 2.5]
)

doc.add_page_break()
add_heading(doc, '3. 功能需求总览', 1)
add_table(doc,
    ['模块', '功能点', '优先级', '说明'],
    [
        ['考核体系管理', '创建考核体系', 'P0', '上传Excel模版，填写体系基本信息'],
        ['考核体系管理', '查看/编辑/删除体系', 'P0', '支持体系的增删改查'],
        ['月度监测', '发起月度监测', 'P0', '选择体系+月份，同一月份同一体系唯一'],
        ['月度监测', '数据收集（整体上传）', 'P0', '管理员上传数据收集页Excel'],
        ['月度监测', '数据收集（分发任务）', 'P0', '根据收数人自动分发，Web填写'],
        ['月度监测', '数据落库', 'P0', '异步生成Excel文件，数据写入数据库'],
        ['月度监测', '确认与发布', 'P0', '机构负责人确认，管理员发布'],
        ['结果展示', '绩效报表', 'P0', 'Excel导出+可视化图表，筛选器支持'],
        ['结果展示', '总览对比报表', 'P0', '分组内机构对比与排名'],
        ['通知', 'IM消息通知', 'P1', '钉钉/飞书/企微动态适配'],
        ['通知', '邮件通知', 'P1', '邮件通知'],
        ['通知', '站内通知', 'P1', '站内消息中心'],
        ['权限', '用户认证与权限', 'P0', '基于角色的权限控制'],
    ],
    col_widths=[2.5, 3.5, 1.5, 5]
)

doc.add_page_break()
add_heading(doc, '4. 考核体系管理', 1)

add_heading(doc, '4.1 Excel模版结构', 2)
doc.add_paragraph(
    '用户上传的Excel模版必须包含4个Sheet页，顺序和表头严格固定。'
    '系统需校验模版格式是否符合规范，不符合则拒绝上传并给出明确提示。'
)

doc.add_heading('4.1.1 模版页（Sheet名：模版页）', 2)
doc.add_paragraph('列顺序严格固定，如下：')
add_table(doc,
    ['序号', '列名', '说明', '数据类型'],
    [
        ['1', '维度', '指标所属维度，如"业务发展""风险控制"', '文本'],
        ['2', '类别', '指标分类，如"存款类""贷款类"', '文本'],
        ['3', '一级指标', '一级指标名称', '文本'],
        ['4', '二级指标', '二级指标名称', '文本'],
        ['5', '权重', '指标权重值，百分比或小数', '数值'],
        ['6', '单位', '指标计量单位', '文本'],
        ['7', '全年目标', '全年目标值', '数值'],
        ['8', '进度目标', '月度进度目标值', '数值'],
        ['9', '实际值', '实际完成值', '数值'],
        ['10', '全年完成率', '实际值/全年目标', '百分比'],
        ['11', '进度完成率', '实际值/进度目标', '百分比'],
        ['12', '指标百分制得分', '由Excel公式计算', '数值'],
        ['13', '指标权重得分', '指标百分制得分×权重', '数值'],
        ['14', '类别得分', '类别内指标权重得分之和', '数值'],
        ['15', '维度得分', '维度内各类别得分之和', '数值'],
        ['16', '总得分', '所有维度得分之和', '数值'],
    ],
    col_widths=[1, 2.5, 5, 2]
)

add_heading(doc, '4.1.2 机构页（Sheet名：机构页）', 2)
add_table(doc,
    ['列名', '数据类型', '说明'],
    [
        ['机构名称', '文本', '分行名称'],
        ['机构ID', '文本', '分行唯一标识'],
        ['分组名称', '文本', '用于报表分组对比，如"东部""西部"'],
        ['机构负责人', '文本', '格式固定为"姓名/工号"，如"张三/EMP001"'],
    ],
    col_widths=[2.5, 2.5, 6]
)

add_heading(doc, '4.1.3 数据收集页（Sheet名：数据收集页）', 2)
doc.add_paragraph(
    '表头占3行，第1行为收数指标名称，第2行为收数人（格式：姓名/工号），'
    '第3行为指标单位。第一列为机构名称，构成机构×指标的二维表。'
)

add_heading(doc, '4.1.4 参数页（Sheet名：参数页）', 2)
add_table(doc,
    ['列名', '说明'],
    [
        ['参数名称', '系统变量名，如 CURRENT_DATE、CURRENT_ORG'],
        ['参数值', '变量值，在生成文件时动态注入'],
    ],
    col_widths=[3, 8]
)

doc.add_heading('4.2 模版工作原理', 2)
doc.add_paragraph(
    '用户上传Excel模版后，系统校验格式并保存。'
    '在月度数据收集阶段，根据模版中的机构页列表和参数页定义，'
    '系统批量复制生成每个机构×月份对应的Excel文件，'
    '同时将当前月份和机构名称注入参数页。'
    '各文件的模版页中通过VLOOKUP等公式引用数据收集页中的数据。'
    '整个过程中原始模版文件不会被修改。'
)

doc.add_page_break()
add_heading(doc, '5. 月度监测流程', 1)

add_heading(doc, '5.1 流程状态', 2)
states = [
    ('待发起', '管理员选择了体系+月份，发起月度监测流程'),
    ('收数中', '收数人填写数据，管理员可整体上传数据收集页'),
    ('收数截止', '到达截止时间或管理员手动截止，收数入口关闭'),
    ('数据落库中', '异步任务：批量生成Excel文件并写入数据库'),
    ('待确认', '各机构负责人分别确认本机构数据'),
    ('已发布', '管理员发布结果，数据公开，机构负责人可查看所有报表'),
]
add_table(doc, ['状态', '说明'], states, col_widths=[2.5, 8])

add_heading(doc, '5.2 流程详细说明', 2)

add_heading(doc, '5.2.1 发起监测', 3)
doc.add_paragraph(
    '管理员选择考核体系和月份，发起月度监测。'
    '系统检查同一体系同一月份是否已有进行中的监测，若有则拒绝。'
    '发起时可配置是否需要审批：若需要，则收数人提交后需管理员审批；若不需要，则直接入库。'
)

add_heading(doc, '5.2.2 数据收集', 3)
doc.add_paragraph('有两种数据收集方式：')
doc.add_paragraph(
    '方式一（整体上传）：管理员直接上传一份完整的数据收集页Excel，'
    '系统自动解析并分发到各机构对应的文件中。'
)
doc.add_paragraph(
    '方式二（任务分发）：系统根据数据收集页中的收数人信息，'
    '将任务自动分发给各收数人。收数人在Web页面填写所有自己负责的指标，'
    '一次性提交。收数人也可下载仅含数据收集页的Excel，填写后上传。'
)

add_heading(doc, '5.2.3 截止与提醒', 3)
doc.add_paragraph(
    '管理员设置收数截止时间。到期前30分钟，系统自动向未提交数据的收数人发送提醒通知。'
    '到达截止时间后，收数入口自动关闭。管理员也可手动提前截止。'
)

add_heading(doc, '5.2.4 数据落库（异步）', 3)
doc.add_paragraph(
    '数据收集截止后，系统启动异步任务：'
)
steps = [
    '根据模版中的机构列表，批量为每个机构×月份生成独立的Excel文件',
    '文件路径：{体系ID}/{年份}/{月份}/{机构ID}/report_{年月}.xlsx，写入MinIO',
    '原始模版不被修改，所有操作基于复制文件进行',
    '从每个生成文件的模版页中读取指标计算结果（实际值、完成率、得分等），写入数据库',
    '同时将数据收集页中的原始数据也落库保存',
    '该过程耗时较长（可能超过1分钟），通过定时任务轮询更新状态，前端展示进度',
]
for s in steps:
    doc.add_paragraph(s, style='List Bullet')

add_heading(doc, '5.2.5 确认与发布', 3)
doc.add_paragraph(
    '数据落库完成后，各机构负责人分别登录系统，确认本机构数据无误后提交确认。'
    '发布前，各负责人只能看到本机构数据，无法查看其他机构。'
    '所有负责人确认后（无需全部同时，逐一确认即可），管理员手动发布。'
    '发布后数据公开，所有机构负责人可查看全部报表并下载文件。'
)

doc.add_page_break()
add_heading(doc, '6. 结果展示', 1)

add_heading(doc, '6.1 绩效报表', 2)
doc.add_paragraph(
    '绩效报表支持Excel导出和可视化图表展示。'
    '筛选条件包括：考核体系、月份、机构。'
    '分组内的机构自动进行对比和排名展示。'
)

add_heading(doc, '6.2 总览对比报表', 2)
doc.add_paragraph(
    '总览报表综合展示某一体系某一月份下所有机构的绩效对比情况，'
    '包括分组内的横向排名、各维度得分雷达图、关键指标柱状图等。'
    '支持按分组维度筛选查看。'
)

doc.add_page_break()
add_heading(doc, '7. 通知模块', 1)

add_heading(doc, '7.1 通知触发场景', 2)
add_table(doc,
    ['场景', '通知方式', '通知对象'],
    [
        ['收数任务分发', 'IM + 邮件 + 站内', '各收数人'],
        ['收数截止前30分钟提醒', 'IM + 邮件 + 站内', '未提交收数人'],
        ['机构负责人待确认提醒', 'IM + 邮件 + 站内', '各机构负责人'],
        ['结果发布通知', 'IM + 邮件 + 站内', '所有相关人员'],
    ],
    col_widths=[3.5, 3, 3]
)

add_heading(doc, '7.2 IM工具适配', 2)
doc.add_paragraph(
    'IM通知模块设计为可插拔架构，通过配置动态选择目标IM工具（钉钉、飞书、企业微信）。'
    '系统通过调用各IM工具的开放API主动推送消息。'
    '具体接入哪种IM工具在系统配置中设置。'
)

doc.add_page_break()
add_heading(doc, '8. 文件存储', 1)

add_heading(doc, '8.1 MinIO路径规范', 2)
add_table(doc,
    ['文件类型', '存储路径', '说明'],
    [
        ['原始模版', 'templates/{体系ID}/template.xlsx', '用户上传的原始模版，不修改'],
        ['月度报告文件', '{体系ID}/{年份}/{月份}/{机构ID}/report_{年月}.xlsx', '每个机构每月一份'],
        ['数据收集页备份', 'uploads/{体系ID}/{年份}/{月份}/{机构ID}/data_collection.xlsx', '每次上传的备份'],
    ],
    col_widths=[3, 6, 3]
)

doc.add_page_break()
add_heading(doc, '9. 非功能需求', 1)

add_heading(doc, '9.1 性能需求', 2)
add_table(doc,
    ['指标', '要求'],
    [
        ['数据落库（100个机构）', '完成时间 ≤ 2分钟'],
        ['Excel上传解析', '单文件 ≤ 5MB，处理时间 ≤ 10秒'],
        ['页面响应时间', '普通操作 ≤ 2秒'],
        ['系统并发用户数', '支持 ≥ 100人同时在线'],
    ],
    col_widths=[5, 6]
)

add_heading(doc, '9.2 安全需求', 2)
items = [
    '所有用户需登录认证，支持用户名密码方式',
    '基于角色的访问控制（RBAC），菜单和数据按角色隔离',
    '机构负责人在发布前只能查看本机构数据',
    '敏感操作（发布、删除）需有操作日志',
    '文件上传需校验文件类型和大小',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

add_heading(doc, '9.3 数据需求', 2)
items = [
    '所有历史数据需持久化保存',
    '同一机构同一月份重复发起监测时系统拒绝',
    '同一机构同一月份重复上传数据时覆盖',
    '原始模版在任何操作中不被修改',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()
add_heading(doc, '附录A：术语表', 1)
add_table(doc,
    ['术语', '说明'],
    [
        ['考核体系', '一套完整的绩效考核方案，包含模版、机构列表和指标定义'],
        ['月度监测', '在某个月份对体系中的指标进行数据收集和评估的过程'],
        ['模版页', 'Excel中记录指标列表及计算公式的Sheet'],
        ['机构页', 'Excel中记录被评价机构信息的Sheet'],
        ['数据收集页', 'Excel中用于填写和分发数据的Sheet'],
        ['参数页', 'Excel中存储系统变量的Sheet，用于文件生成时注入'],
        ['收数人', '被指定负责填写某些指标数据的人员'],
        ['机构负责人', '分行领导，有权限查看和确认本机构绩效数据'],
    ],
    col_widths=[3, 8]
)

doc.save('/Users/zhaoyu/Desktop/机构绩效管理系统_产品需求文档.docx')
print('PRD generated successfully')
