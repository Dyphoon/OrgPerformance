#!/usr/bin/env python3
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

def set_cell_bg(cell, color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)

def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        set_cell_bg(cell, '2E4057')
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(10)
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        bg = 'F8F9FA' if ri % 2 == 0 else 'FFFFFF'
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = str(val)
            set_cell_bg(cell, bg)
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.font.size = Pt(9)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    return table

doc = Document()
section = doc.sections[0]
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)

title = doc.add_heading('机构绩效管理系统\n技术设计文档', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph(f'版本：V1.0.0\n日期：{datetime.date.today().strftime("%Y-%m-%d")}\n状态：初稿').alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_page_break()

add_heading(doc, '目录', 1)
doc.add_paragraph('1. 技术架构概述\n2. 技术选型说明\n3. 系统架构设计\n4. 数据库设计\n5. 核心模块设计\n6. API接口设计\n7. MinIO文件存储设计\n8. 异步任务设计\n9. 通知模块设计\n10. 前端架构设计\n11. 安全设计\n12. 部署方案')
doc.add_page_break()

add_heading(doc, '1. 技术架构概述', 1)
doc.add_paragraph('本系统采用前后分离的B/S架构，后端基于Java Spring Boot，前端基于React。系统分为四个主要模块：考核体系管理、月度监测管理、数据落库处理、结果展示。文件存储采用MinIO对象存储，Excel处理使用Apache POI库。')
doc.add_page_break()

add_heading(doc, '2. 技术选型说明', 1)
add_table(doc, ['技术项', '选型', '版本', '说明'], [
    ['后端框架', 'Spring Boot', '3.x', '主框架，提供RESTful API'],
    ['ORM', 'MyBatis', '3.0', '轻量级ORM，SQL可控'],
    ['数据库', 'MySQL', '8.0', '关系型数据库，存储业务数据'],
    ['文件存储', 'MinIO', '最新', 'S3兼容的对象存储'],
    ['Excel处理', 'Apache POI', '5.x', 'Java Excel处理库'],
    ['前端框架', 'React', '19.x', 'UI框架'],
    ['UI组件库', 'Ant Design', '6.x', '企业级React组件库'],
    ['图表', 'ECharts', '6.x', '可视化图表'],
    ['构建工具', 'Vite', '8.x', '前端构建工具'],
    ['语言', 'TypeScript', '5.9', '前端主语言'],
    ['通知', 'HTTP调用IM开放API', '-', '钉钉/飞书/企微'],
], col_widths=[2.5, 3, 1.5, 4])
doc.add_page_break()

add_heading(doc, '3. 系统架构设计', 1)
add_heading(doc, '3.1 整体架构', 2)
doc.add_paragraph('系统分为五层：')
for name, desc in [
    ('接入层', '用户通过浏览器访问React前端，所有请求经HTTPS到后端服务'),
    ('网关层', 'Spring Boot内置Tomcat作为Web容器，提供RESTful API'),
    ('业务层', '核心业务逻辑，包括体系管理、监测流程、数据处理'),
    ('数据层', 'MySQL存储业务数据，MinIO存储文件'),
    ('任务层', '定时任务处理异步操作（数据落库、通知提醒）'),
]:
    p = doc.add_paragraph()
    p.add_run(name + '：').bold = True
    p.add_run(desc)

add_heading(doc, '3.2 包结构设计（后端）', 2)
add_table(doc, ['包名', '说明'], [
    ['controller', 'REST控制器，接收HTTP请求'],
    ['service', '业务逻辑层'],
    ['mapper', 'MyBatis Mapper接口和XML'],
    ['entity', '数据库实体类'],
    ['dto', '数据传输对象'],
    ['vo', '视图对象'],
    ['job', '定时任务'],
    ['notify', '通知模块（IM/邮件/站内）'],
    ['excel', 'Excel处理相关类'],
    ['minio', 'MinIO文件操作封装'],
    ['config', '系统配置类'],
], col_widths=[3, 8])

add_heading(doc, '3.3 目录结构（前端）', 2)
add_table(doc, ['目录', '说明'], [
    ['src/api', 'API调用封装'],
    ['src/pages', '页面组件，按菜单模块划分'],
    ['src/components', '公共组件'],
    ['src/store', '状态管理'],
    ['src/hooks', '自定义Hooks'],
    ['src/types', 'TypeScript类型定义'],
    ['src/utils', '工具函数'],
], col_widths=[3, 8])
doc.add_page_break()

add_heading(doc, '4. 数据库设计', 1)
add_heading(doc, '4.1 ER图概述', 2)
doc.add_paragraph('核心实体包括：考核体系（AssessmentSystem）、月度监测（MonthlyMonitoring）、机构（Institution）、指标（Indicator）、月度指标数据（MonthlyIndicatorData）、数据收集任务（CollectionTask）、用户（User）、角色（Role）、通知记录（Notification）。')
rels = [
    '考核体系 1:N 月度监测（一个体系可有多个月份的监测）',
    '考核体系 1:N 机构（一个体系包含多个被考核机构）',
    '月度监测 1:N 数据收集任务（一次监测包含多个收数任务）',
    '月度监测 1:N 月度指标数据（一次监测包含多条指标数据）',
    '机构 1:N 数据收集任务（一个机构对应多条收数任务）',
    '用户 N:M 角色（多对多关系）',
]
for r in rels:
    doc.add_paragraph(r, style='List Bullet')

add_heading(doc, '4.2 核心表结构', 2)

add_heading(doc, '4.2.1 考核体系表 assessment_system', 3)
add_table(doc, ['字段', '类型', '说明'], [
    ['id', 'BIGINT PK', '主键'],
    ['name', 'VARCHAR(100)', '体系名称'],
    ['description', 'TEXT', '体系描述'],
    ['template_file_key', 'VARCHAR(500)', '原始模版在MinIO中的路径'],
    ['need_approval', 'TINYINT(1)', '是否需要审批：0-否，1-是'],
    ['status', 'TINYINT', '状态：0-禁用，1-启用'],
    ['created_by', 'VARCHAR(50)', '创建人'],
    ['created_at', 'DATETIME', '创建时间'],
    ['updated_at', 'DATETIME', '更新时间'],
], col_widths=[3, 3, 5])

add_heading(doc, '4.2.2 月度监测表 monthly_monitoring', 3)
add_table(doc, ['字段', '类型', '说明'], [
    ['id', 'BIGINT PK', '主键'],
    ['system_id', 'BIGINT FK', '关联体系ID'],
    ['year', 'INT', '年份'],
    ['month', 'INT', '月份'],
    ['status', 'VARCHAR(20)', '状态：PENDING/collecting/closed/processing/confirming/published'],
    ['deadline', 'DATETIME', '收数截止时间'],
    ['approval_required', 'TINYINT(1)', '是否需要审批（从体系继承+可覆盖）'],
    ['process_percent', 'INT', '数据落库进度百分比'],
    ['process_status', 'VARCHAR(20)', '处理状态：idle/processing/done/failed'],
    ['process_msg', 'VARCHAR(500)', '处理信息或错误描述'],
    ['created_by', 'VARCHAR(50)', '发起人'],
    ['created_at', 'DATETIME', '创建时间'],
    ['updated_at', 'DATETIME', '更新时间'],
], col_widths=[3, 3, 5])

add_heading(doc, '4.2.3 机构表 institution', 3)
add_table(doc, ['字段', '类型', '说明'], [
    ['id', 'BIGINT PK', '主键'],
    ['system_id', 'BIGINT FK', '关联体系ID'],
    ['org_name', 'VARCHAR(100)', '机构名称'],
    ['org_id', 'VARCHAR(50)', '机构ID（来自Excel）'],
    ['group_name', 'VARCHAR(100)', '分组名称'],
    ['leader_name', 'VARCHAR(50)', '负责人姓名'],
    ['leader_emp_no', 'VARCHAR(50)', '负责人工号'],
    ['created_at', 'DATETIME', '创建时间'],
], col_widths=[3, 3, 5])

add_heading(doc, '4.2.4 指标表 indicator', 3)
add_table(doc, ['字段', '类型', '说明'], [
    ['id', 'BIGINT PK', '主键'],
    ['system_id', 'BIGINT FK', '关联体系ID'],
    ['dimension', 'VARCHAR(50)', '维度'],
    ['category', 'VARCHAR(50)', '类别'],
    ['level1_name', 'VARCHAR(100)', '一级指标'],
    ['level2_name', 'VARCHAR(100)', '二级指标'],
    ['weight', 'DECIMAL(10,4)', '权重'],
    ['unit', 'VARCHAR(20)', '单位'],
    ['annual_target', 'DECIMAL(20,4)', '全年目标'],
    ['progress_target', 'DECIMAL(20,4)', '进度目标（月度）'],
    ['row_index', 'INT', '在Excel中的行号'],
    ['created_at', 'DATETIME', '创建时间'],
], col_widths=[3, 3, 5])

add_heading(doc, '4.2.5 月度指标数据表 monthly_indicator_data', 3)
add_table(doc, ['字段', '类型', '说明'], [
    ['id', 'BIGINT PK', '主键'],
    ['monitoring_id', 'BIGINT FK', '关联月度监测ID'],
    ['indicator_id', 'BIGINT FK', '关联指标ID'],
    ['institution_id', 'BIGINT FK', '关联机构ID'],
    ['actual_value', 'DECIMAL(20,4)', '实际值'],
    ['annual_completion_rate', 'DECIMAL(10,4)', '全年完成率'],
    ['progress_completion_rate', 'DECIMAL(10,4)', '进度完成率'],
    ['score_100', 'DECIMAL(10,4)', '指标百分制得分'],
    ['score_weighted', 'DECIMAL(10,4)', '指标权重得分'],
    ['score_category', 'DECIMAL(10,4)', '类别得分'],
    ['score_dimension', 'DECIMAL(10,4)', '维度得分'],
    ['total_score', 'DECIMAL(10,4)', '总得分'],
    ['file_key', 'VARCHAR(500)', '对应生成的Excel在MinIO中的路径'],
    ['created_at', 'DATETIME', '落库时间'],
], col_widths=[3, 3, 5])

add_heading(doc, '4.2.6 数据收集任务表 collection_task', 3)
add_table(doc, ['字段', '类型', '说明'], [
    ['id', 'BIGINT PK', '主键'],
    ['monitoring_id', 'BIGINT FK', '关联月度监测ID'],
    ['indicator_id', 'BIGINT FK', '关联指标ID（若是整体上传则为空）'],
    ['institution_id', 'BIGINT FK', '关联机构ID'],
    ['collector_name', 'VARCHAR(50)', '收数人姓名'],
    ['collector_emp_no', 'VARCHAR(50)', '收数人工号'],
    ['collector_user_id', 'BIGINT FK', '收数人系统用户ID（nullable）'],
    ['actual_value', 'DECIMAL(20,4)', '实际填写的数值（nullable）'],
    ['status', 'VARCHAR(20)', '状态：pending/submitted/approved/rejected'],
    ['submitted_at', 'DATETIME', '提交时间'],
    ['approved_by', 'VARCHAR(50)', '审批人'],
    ['approved_at', 'DATETIME', '审批时间'],
    ['remark', 'TEXT', '备注'],
    ['created_at', 'DATETIME', '创建时间'],
], col_widths=[3, 3, 5])

add_heading(doc, '4.2.7 用户表 sys_user', 3)
add_table(doc, ['字段', '类型', '说明'], [
    ['id', 'BIGINT PK', '主键'],
    ['username', 'VARCHAR(50)', '用户名（工号）'],
    ['password', 'VARCHAR(200)', '密码（加密存储）'],
    ['name', 'VARCHAR(50)', '姓名'],
    ['emp_no', 'VARCHAR(50)', '工号'],
    ['email', 'VARCHAR(100)', '邮箱'],
    ['phone', 'VARCHAR(20)', '手机号'],
    ['status', 'TINYINT', '状态：0-禁用，1-启用'],
    ['created_at', 'DATETIME', '创建时间'],
    ['updated_at', 'DATETIME', '更新时间'],
], col_widths=[3, 3, 5])

add_heading(doc, '4.2.8 角色表 sys_role', 3)
add_table(doc, ['字段', '类型', '说明'], [
    ['id', 'BIGINT PK', '主键'],
    ['role_code', 'VARCHAR(50)', '角色代码：admin/collector/leader'],
    ['role_name', 'VARCHAR(50)', '角色名称'],
    ['description', 'VARCHAR(200)', '角色描述'],
], col_widths=[3, 3, 5])

add_heading(doc, '4.2.9 用户角色表 sys_user_role', 3)
add_table(doc, ['字段', '类型', '说明'], [
    ['user_id', 'BIGINT FK', '用户ID'],
    ['role_id', 'BIGINT FK', '角色ID'],
], col_widths=[3, 3])

add_heading(doc, '4.2.10 通知记录表 notification', 3)
add_table(doc, ['字段', '类型', '说明'], [
    ['id', 'BIGINT PK', '主键'],
    ['user_id', 'BIGINT FK', '通知对象用户ID'],
    ['title', 'VARCHAR(200)', '通知标题'],
    ['content', 'TEXT', '通知内容'],
    ['type', 'VARCHAR(20)', '类型：im/email/site'],
    ['status', 'VARCHAR(20)', '状态：pending/sent/failed'],
    ['send_result', 'VARCHAR(500)', '发送结果或错误信息'],
    ['sent_at', 'DATETIME', '发送时间'],
    ['created_at', 'DATETIME', '创建时间'],
], col_widths=[3, 3, 5])

add_heading(doc, '4.2.11 机构负责人关联表 institution_leader', 3)
add_table(doc, ['字段', '类型', '说明'], [
    ['institution_id', 'BIGINT FK', '机构ID'],
    ['user_id', 'BIGINT FK', '负责人用户ID'],
    ['confirmed', 'TINYINT(1)', '是否已确认：0-否，1-是'],
    ['confirmed_at', 'DATETIME', '确认时间'],
], col_widths=[3, 3, 5])

add_heading(doc, '4.2.12 确认记录表 confirmation_record', 3)
add_table(doc, ['字段', '类型', '说明'], [
    ['id', 'BIGINT PK', '主键'],
    ['monitoring_id', 'BIGINT FK', '关联月度监测ID'],
    ['institution_id', 'BIGINT FK', '关联机构ID'],
    ['user_id', 'BIGINT FK', '确认人用户ID'],
    ['status', 'VARCHAR(20)', '状态：pending/confirmed'],
    ['confirmed_at', 'DATETIME', '确认时间'],
    ['remark', 'VARCHAR(500)', '确认备注'],
], col_widths=[3, 3, 5])
doc.add_page_break()

add_heading(doc, '5. 核心模块设计', 1)
add_heading(doc, '5.1 考核体系管理', 2)
for f in [
    '上传Excel模版：校验4个Sheet格式，解析机构列表和指标列表，保存原始文件到MinIO',
    '创建体系：将模版解析结果写入数据库（机构表、指标表）',
    '编辑/删除体系：删除时同时删除关联的监测记录和MinIO文件',
    '体系列表：分页展示所有体系，支持按名称搜索',
]:
    doc.add_paragraph(f, style='List Bullet')

add_heading(doc, '5.2 月度监测流程', 2)
add_table(doc, ['状态', '可执行操作', '后端处理'], [
    ['PENDING', '发起监测', '创建月度监测记录，状态变为COLLECTING'],
    ['COLLECTING', '整体上传/任务分发/手动截止', '整体上传解析数据写入任务表；任务分发创建任务记录'],
    ['COLLECTING->CLOSED', '截止时间到达或手动截止', '状态变为CLOSED，关闭收数入口'],
    ['CLOSED->PROCESSING', '自动触发（定时任务）', '启动异步数据落库流程'],
    ['PROCESSING', '-', '异步生成Excel+落库，通过轮询更新进度'],
    ['PROCESSING->CONFIRMING', '数据落库完成', '状态变为CONFIRMING，通知各机构负责人'],
    ['CONFIRMING', '负责人确认', '各负责人提交确认记录'],
    ['CONFIRMING->PUBLISHED', '管理员发布', '状态变为PUBLISHED，数据公开'],
], col_widths=[2.5, 3.5, 5])

add_heading(doc, '5.3 数据落库模块', 2)
doc.add_paragraph('异步处理流程（定时任务轮询）：')
add_table(doc, ['步骤', '操作', '说明'], [
    ['Step 1', '扫描状态为CLOSED且process_status=idle的监测记录', '避免重复触发'],
    ['Step 2', '将监测记录process_status更新为processing', '加锁防止并发'],
    ['Step 3', '获取该监测关联的所有机构和指标', '从数据库查询'],
    ['Step 4', '遍历机构列表，为每个机构×月份生成Excel文件', '复制模版，注入参数页'],
    ['Step 5', '将生成的Excel上传到MinIO，记录file_key', '按路径规范存储'],
    ['Step 6', '读取生成文件的模版页，解析计算结果', '使用POI读取公式计算后的值'],
    ['Step 7', '将模版页数据写入monthly_indicator_data表', '每条指标一行记录'],
    ['Step 8', '将数据收集页原始数据写入collection_task表', '保留收数原始记录'],
    ['Step 9', '更新process_status=done，process_percent=100', '标记完成'],
    ['Step 10', '定时任务每10秒轮询一次，更新前端进度', 'process_percent实时更新'],
], col_widths=[2, 5, 4])

add_heading(doc, '5.4 确认发布模块', 2)
doc.add_paragraph('机构负责人登录后，系统查询其关联的机构及进行中的CONFIRMING状态监测，展示本机构数据。负责人提交确认后，记录确认信息。管理员查看确认进度，确认全部完成后手动发布。发布后，系统将监测状态更新为PUBLISHED，同时向所有相关人员发送发布通知。')
doc.add_page_break()

add_heading(doc, '6. API接口设计', 1)
add_heading(doc, '6.1 认证接口', 2)
add_table(doc, ['接口', '方法', '说明'], [
    ['/api/auth/login', 'POST', '登录，返回JWT Token'],
    ['/api/auth/logout', 'POST', '登出'],
    ['/api/auth/current-user', 'GET', '获取当前用户信息'],
], col_widths=[4, 1.5, 5])

add_heading(doc, '6.2 考核体系接口', 2)
add_table(doc, ['接口', '方法', '说明'], [
    ['/api/systems', 'GET', '体系列表（分页+搜索）'],
    ['/api/systems', 'POST', '创建体系（上传模版）'],
    ['/api/systems/{id}', 'GET', '体系详情（含机构、指标）'],
    ['/api/systems/{id}', 'PUT', '编辑体系'],
    ['/api/systems/{id}', 'DELETE', '删除体系'],
    ['/api/systems/{id}/template/download', 'GET', '下载原始模版'],
], col_widths=[4, 1.5, 5])

add_heading(doc, '6.3 月度监测接口', 2)
add_table(doc, ['接口', '方法', '说明'], [
    ['/api/monitorings', 'GET', '监测列表（分页+筛选）'],
    ['/api/monitorings', 'POST', '发起月度监测'],
    ['/api/monitorings/{id}', 'GET', '监测详情'],
    ['/api/monitorings/{id}/collect/upload', 'POST', '整体上传数据收集页Excel'],
    ['/api/monitorings/{id}/collect/tasks', 'GET', '获取收数任务列表'],
    ['/api/monitorings/{id}/collect/tasks/{taskId}', 'PUT', '收数人提交任务数据'],
    ['/api/monitorings/{id}/close', 'POST', '手动截止收数'],
    ['/api/monitorings/{id}/process/status', 'GET', '查询数据落库进度（轮询）'],
    ['/api/monitorings/{id}/confirm', 'POST', '机构负责人确认'],
    ['/api/monitorings/{id}/publish', 'POST', '管理员发布'],
], col_widths=[5, 1.5, 5])

add_heading(doc, '6.4 结果展示接口', 2)
add_table(doc, ['接口', '方法', '说明'], [
    ['/api/reports/summary', 'GET', '总览报表（所有机构对比）'],
    ['/api/reports/institution/{id}', 'GET', '特定机构报表'],
    ['/api/reports/export/excel', 'GET', '导出Excel报表'],
    ['/api/reports/download/{monitoringId}/{institutionId}', 'GET', '下载机构月度报告Excel'],
], col_widths=[5, 1.5, 5])

add_heading(doc, '6.5 通知接口', 2)
add_table(doc, ['接口', '方法', '说明'], [
    ['/api/notifications', 'GET', '站内通知列表（当前用户）'],
    ['/api/notifications/{id}/read', 'PUT', '标记已读'],
    ['/api/notifications/unread-count', 'GET', '未读通知数量'],
], col_widths=[5, 1.5, 5])
doc.add_page_break()

add_heading(doc, '7. MinIO文件存储设计', 1)
add_heading(doc, '7.1 Bucket设计', 2)
add_table(doc, ['Bucket', '用途', '访问策略'], [
    ['perf-templates', '存储原始模版文件', 'Private'],
    ['perf-reports', '存储生成的月度报告文件', 'Private（发布后提供下载）'],
    ['perf-uploads', '存储用户上传的数据收集页备份', 'Private'],
], col_widths=[3.5, 5, 3])

add_heading(doc, '7.2 路径规范', 2)
add_table(doc, ['文件类型', 'MinIO路径', '说明'], [
    ['原始模版', 'perf-templates/{体系ID}/template_{timestamp}.xlsx', '上传时带时间戳区分版本'],
    ['月度报告', 'perf-reports/{体系ID}/{年份}/{月份}/{机构ID}/report_{年月}.xlsx', '按体系/年月/机构分层'],
    ['数据收集页备份', 'perf-uploads/{体系ID}/{年份}/{月份}/{机构ID}/data_{timestamp}.xlsx', '每次上传独立文件'],
], col_widths=[3, 5, 3])

add_heading(doc, '7.3 访问控制', 2)
doc.add_paragraph('MinIO所有文件默认私有（Private）。发布后的报告文件通过后端接口验证权限后，生成预签名URL（Presigned URL）提供给用户下载，有效期设为15分钟。未发布的文件即使知道路径也无法直接访问。')
doc.add_page_break()

add_heading(doc, '8. 异步任务设计', 1)
add_heading(doc, '8.1 定时任务清单', 2)
add_table(doc, ['任务', 'cron表达式', '说明'], [
    ['DataProcessingJob（数据落库轮询）', '*/10 * * * * *', '每10秒扫描CLOSED状态监测，触发数据落库'],
    ['DeadlineReminderJob（截止提醒）', '0 * * * * *', '每小时检查收数截止时间，发起前30分钟提醒'],
    ['MonitoringStatusSyncJob（状态同步）', '*/30 * * * * *', '每30秒同步监测process_percent到数据库'],
], col_widths=[4, 3, 4])

add_heading(doc, '8.2 数据落库任务状态机', 2)
doc.add_paragraph('process_status字段控制任务状态：idle -> processing -> done/failed。每次扫描到CLOSED且idle的记录，先将status改为processing，然后执行业务逻辑，完成后改为done，失败改为failed并记录错误信息。process_percent字段实时更新进度百分比，前端通过轮询API获取。')
doc.add_page_break()

add_heading(doc, '9. 通知模块设计', 1)
add_heading(doc, '9.1 通知服务架构', 2)
doc.add_paragraph('通知服务采用策略模式设计，核心接口 NotifyStrategy，包含方法：send(title, content, recipient)。具体实现：')
for name, desc in [
    ('DingTalkNotifyStrategy', '调用钉钉自定义机器人Webhook API'),
    ('FeishuNotifyStrategy', '调用飞书自定义机器人Webhook API'),
    ('WechatWorkNotifyStrategy', '调用企业微信群机器人Webhook API'),
    ('EmailNotifyStrategy', '调用SMTP服务发送邮件'),
    ('SiteNotifyStrategy', '写入notification表，前端轮询拉取'),
]:
    p = doc.add_paragraph()
    p.add_run(name + '：').bold = True
    p.add_run(desc)

add_heading(doc, '9.2 通知触发时机', 2)
add_table(doc, ['触发时机', '通知对象', '内容摘要'], [
    ['发起监测', '所有收数人', '您有新的收数任务待填写，体系：{体系名}，月份：{年月}'],
    ['截止前30分钟', '未提交收数人', '收数即将截止，请尽快提交'],
    ['数据落库完成', '所有机构负责人', '本月绩效数据已生成，请确认'],
    ['结果发布', '所有相关人员', '绩效结果已发布，请查看'],
], col_widths=[3, 3, 5])

add_heading(doc, '9.3 配置管理', 2)
doc.add_paragraph('在系统配置表（sys_config）中存储各IM工具的webhook地址和密钥，支持动态切换目标IM工具，无需修改代码。邮件配置使用Spring Mail，配置SMTP服务器信息。')
doc.add_page_break()

add_heading(doc, '10. 前端架构设计', 1)
add_heading(doc, '10.1 路由设计', 2)
add_table(doc, ['路由', '页面', '权限'], [
    ['/login', '登录页', '公开'],
    ['/dashboard', '首页/总览', '管理员+机构负责人'],
    ['/system', '考核体系管理', '管理员'],
    ['/monitoring', '月度监测列表', '管理员'],
    ['/monitoring/:id', '监测详情', '管理员'],
    ['/collect', '数据填写（收数人）', '收数人'],
    ['/report', '绩效报表', '管理员+机构负责人（发布后）'],
    ['/notifications', '通知中心', '所有登录用户'],
], col_widths=[3.5, 4, 3])

add_heading(doc, '10.2 核心页面组件', 2)
add_table(doc, ['组件', '说明'], [
    ['SystemListPage', '考核体系列表页，含搜索、分页、创建入口'],
    ['SystemCreatePage', '创建体系页面，含Excel上传、模版预览'],
    ['MonitoringListPage', '月度监测列表页，展示各状态下的监测'],
    ['MonitoringDetailPage', '监测详情页，含数据收集、进度、确认状态'],
    ['DataCollectPage', '数据填写页（收数人），展示自己负责的指标表单'],
    ['ReportPage', '报表展示页，含筛选器、ECharts图表'],
    ['OverviewPage', '总览对比页，分组机构排名、雷达图、柱状图'],
    ['NotificationPage', '通知中心页，展示所有通知'],
], col_widths=[3, 8])

add_heading(doc, '10.3 状态管理', 2)
doc.add_paragraph('使用React Context管理全局状态：AuthContext（认证信息）、NotificationContext（未读通知数）、MonitoringContext（当前监测数据）。表单数据使用React Hook Form管理，校验使用Zod。')

add_heading(doc, '10.4 数据展示', 2)
doc.add_paragraph('绩效报表使用ECharts实现，核心图表类型：分组机构排名使用柱状图（横向），多维度得分对比使用雷达图，全年进度使用折线图，指标权重分布使用饼图。所有图表支持按机构、月份、体系动态切换数据。')
doc.add_page_break()

add_heading(doc, '11. 安全设计', 1)
add_heading(doc, '11.1 认证', 2)
doc.add_paragraph('使用JWT（JSON Web Token）实现无状态认证。用户登录成功后，后端签发JWT Token，有效期24小时。前端在请求头中携带Token（Authorization: Bearer {token}）。后端通过Spring Security Filter验证Token有效性。')

add_heading(doc, '11.2 授权', 2)
add_table(doc, ['接口类型', '权限要求'], [
    ['考核体系CRUD', '管理员角色'],
    ['发起监测/截止/发布', '管理员角色'],
    ['数据填写/提交', '收数人角色（且必须是该任务的收数人）'],
    ['机构负责人确认', '机构负责人角色（且必须是该机构的负责人）'],
    ['绩效报表查看', '管理员（所有）、机构负责人（本机构发布前，所有机构发布后）'],
], col_widths=[5, 6])

add_heading(doc, '11.3 数据隔离', 2)
doc.add_paragraph('机构负责人在CONFIRMING状态下，只能查询到自己关联机构的数据。PUBLISHED状态后，机构负责人可查询所有机构的数据用于对比。收数人只能查询和填写自己被分配的指标数据。所有数据查询必须带institution_id和monitoring_id的关联校验。')

add_heading(doc, '11.4 文件安全', 2)
doc.add_paragraph('Excel上传限制：文件大小不超过5MB，后缀必须为.xlsx。文件内容通过POI解析，禁止执行宏。所有文件通过MinIO存储，不直接暴露文件路径。文件下载通过后端接口鉴权后生成预签名URL。')
doc.add_page_break()

add_heading(doc, '12. 部署方案', 1)
add_heading(doc, '12.1 环境规划', 2)
add_table(doc, ['环境', '用途', '部署组件'], [
    ['开发环境', '本地开发', '后端JAR + 前端Vite DevServer + MySQL (all in one)'],
    ['测试环境', '功能测试', '后端JAR + 前端Nginx + MySQL + MinIO'],
    ['生产环境', '正式使用', '后端JAR（多实例） + 前端Nginx + MySQL + MinIO'],
], col_widths=[2.5, 3, 5.5])

add_heading(doc, '12.2 配置文件', 2)
add_table(doc, ['配置项', '说明'], [
    ['spring.datasource.url', 'MySQL连接地址'],
    ['spring.minio.endpoint', 'MinIO服务地址'],
    ['spring.minio.access-key / secret-key', 'MinIO访问凭证'],
    ['notify.im.type', 'IM工具类型（dingtalk/feishu/wechatwork）'],
    ['notify.im.webhook.*', '各IM工具的webhook地址'],
    ['spring.mail.*', '邮件SMTP配置'],
    ['jwt.secret', 'JWT签名密钥'],
], col_widths=[4, 7])

add_heading(doc, '12.3 构建命令', 2)
doc.add_paragraph('后端构建：./mvnw clean package -DskipTests')
doc.add_paragraph('前端构建：cd frontend && npm install && npm run build')
doc.add_page_break()

add_heading(doc, '附录：数据库ER关系描述', 1)
doc.add_paragraph('assessment_system（考核体系） 1——N monthly_monitoring（月度监测）\nassessment_system 1——N institution（机构）\nassessment_system 1——N indicator（指标）\nmonthly_monitoring 1——N monthly_indicator_data（月度指标数据）\nmonthly_monitoring 1——N collection_task（数据收集任务）\nmonthly_monitoring 1——N confirmation_record（确认记录）\ninstitution 1——N institution_leader（机构负责人关联）\ninstitution_leader N——1 sys_user（用户）\ncollection_task N——1 sys_user（收数人用户）\nindicator N——1 institution（指标归属机构）\nmonthly_indicator_data N——1 indicator\nmonthly_indicator_data N——1 institution\nsys_user N——M sys_role（用户角色）\nnotification N——1 sys_user')

doc.save('/Users/zhaoyu/Desktop/机构绩效管理系统_技术设计文档.docx')
print('Technical design doc generated successfully')
