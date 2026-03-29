#!/usr/bin/env python3
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

def set_cell_bg(cell, color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)

def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)

def add_code_block(doc, code_text):
    p = doc.add_paragraph()
    run = p.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(8)
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        set_cell_bg(cell, '2E4057')
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(9)
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        bg = 'F8F9FA' if ri % 2 == 0 else 'FFFFFF'
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = str(val)
            set_cell_bg(cell, bg)
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in para.runs:
                    run.font.size = Pt(8)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    return table

def section(doc, title):
    add_heading(doc, title, 1)
    doc.add_page_break()

def api_group(doc, group_name):
    add_heading(doc, group_name, 2)

def api_entry(doc, method, path, desc):
    p = doc.add_paragraph()
    run = p.add_run(f'{method} {path}')
    run.bold = True
    run.font.color.rgb = RGBColor(0, 82, 178) if method == 'GET' else RGBColor(0, 118, 56) if method == 'POST' else RGBColor(163, 21, 21)
    p.add_run(f'  {desc}')

def req_header(doc, headers):
    add_heading(doc, '请求头', 4)
    add_table(doc, ['参数', '类型', '必填', '说明'], headers, col_widths=[3, 2, 1.5, 5])

def req_param(doc, params):
    add_heading(doc, '请求参数', 4)
    if params:
        add_table(doc, ['参数名', '位置', '类型', '必填', '说明'], params, col_widths=[2.5, 1.5, 2, 1.5, 4])
    else:
        doc.add_paragraph('无')

def req_body(doc, body):
    add_heading(doc, '请求体（JSON）', 4)
    if body:
        add_table(doc, ['字段', '类型', '必填', '说明'], body, col_widths=[3, 2, 1.5, 5])
    else:
        doc.add_paragraph('无')

def req_file(doc):
    add_heading(doc, '请求体（multipart/form-data）', 4)
    add_table(doc, ['字段', '类型', '必填', '说明'], [
        ['file', 'File', '是', 'Excel文件，后缀.xlsx，大小<=5MB'],
    ], col_widths=[2.5, 2, 1.5, 5])

def resp(doc, code, desc):
    add_heading(doc, '响应示例', 4)
    add_heading(doc, f'HTTP {code} {desc}', 5)
    doc.add_paragraph('')

def resp_success(doc, json_example, desc='成功响应'):
    add_heading(doc, desc, 4)
    code_block = json_example
    p = doc.add_paragraph()
    run = p.add_run(code_block)
    run.font.name = 'Courier New'
    run.font.size = Pt(8)
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)

def resp_error(doc, errors):
    add_heading(doc, '错误响应', 4)
    add_table(doc, ['HTTP状态码', '错误码', '说明'], errors, col_widths=[3, 3, 5])

doc = Document()
section = doc.sections[0]
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)

title = doc.add_heading('机构绩效管理系统\nAPI接口详细设计', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph(f'版本：V1.0.0\n日期：{datetime.date.today().strftime("%Y-%m-%d")}\n状态：初稿').alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_page_break()

add_heading(doc, '目录', 1)
doc.add_paragraph(
    '1. 接口规范\n'
    '2. 认证相关接口\n'
    '3. 考核体系接口\n'
    '4. 月度监测接口\n'
    '5. 数据收集接口\n'
    '6. 结果展示接口\n'
    '7. 通知接口\n'
    '8. 仪表盘接口'
)

doc.add_page_break()
add_heading(doc, '1. 接口规范', 1)

add_heading(doc, '1.1 基础规范', 2)
add_table(doc, ['项目', '规范'], [
    ['基础URL', '/api'],
    ['认证方式', 'JWT Bearer Token，放在请求头 Authorization: Bearer {token}'],
    ['Content-Type', 'application/json（除文件上传接口使用 multipart/form-data）'],
    ['字符编码', 'UTF-8'],
    ['时间格式', 'ISO 8601格式：yyyy-MM-ddTHH:mm:ss.SSSZ，如 "2026-03-27T10:30:00.000+08:00"'],
    ['日期格式', 'yyyy-MM-dd，如 "2026-03-27"'],
    ['数字精度', '小数统一保留4位（百分制得分等），百分比显示时乘100'],
    ['分页参数', 'page=1（从1开始），pageSize=10，max pageSize=100'],
    ['排序参数', 'sort=createdAt,desc（格式：字段,asc|desc）'],
], col_widths=[3, 8])

add_heading(doc, '1.2 统一响应结构', 2)
resp_success(doc,
'{\n'
'  "code": 200,          // 业务状态码，200=成功，非200=失败\n'
'  "message": "success", // 描述信息\n'
'  "data": {             // 数据体，成功时有值，失败时为null\n'
'    "id": 1,\n'
'    "name": "xxx"\n'
'  },\n'
'  "timestamp": "2026-03-27T10:30:00.000+08:00",\n'
'  "requestId": "uuid-xxx"  // 请求追踪ID\n'
'}',
'成功响应结构')

add_heading(doc, '1.3 分页响应结构', 3)
resp_success(doc,
'{\n'
'  "code": 200,\n'
'  "message": "success",\n'
'  "data": {\n'
'    "records": [],       // 数据列表\n'
'    "total": 100,        // 总记录数\n'
'    "page": 1,           // 当前页\n'
'    "pageSize": 10,      // 每页大小\n'
'    "pages": 10          // 总页数\n'
'  },\n'
'  "timestamp": "2026-03-27T10:30:00.000+08:00"\n'
'}',
'分页响应结构（data内嵌分页信息）')

add_heading(doc, '1.4 错误码定义', 2)
add_table(doc, ['错误码', '说明', '处理建议'], [
    ['200', '操作成功', '正常处理'],
    ['400', '请求参数错误', '检查请求参数格式和必填项'],
    ['401', '未认证或Token过期', '重新登录获取Token'],
    ['403', '无权限访问', '检查当前用户角色和权限'],
    ['404', '资源不存在', '检查请求的资源ID是否正确'],
    ['409', '业务冲突', '如同一月份重复发起监测，返回此错误'],
    ['413', '文件过大', '上传文件超过5MB限制'],
    ['415', '文件格式不支持', '仅支持.xlsx格式'],
    ['500', '服务端内部错误', '联系管理员排查'],
], col_widths=[2, 4, 5])

doc.add_page_break()
add_heading(doc, '2. 认证相关接口', 1)

api_group(doc, '2.1 登录')
api_entry(doc, 'POST', '/api/auth/login', '用户登录，获取JWT Token')
req_body(doc, [
    ['username', 'string', '是', '用户名（工号）'],
    ['password', 'string', '是', '密码（明文传输，生产环境必须HTTPS）'],
])
resp_success(doc,
'{\n'
'  "code": 200,\n'
'  "message": "success",\n'
'  "data": {\n'
'    "token": "eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9...",\n'
'    "expiresIn": 86400,\n'
'    "tokenType": "Bearer",\n'
'    "user": {\n'
'      "id": 1,\n'
'      "username": "EMP001",\n'
'      "name": "张三",\n'
'      "empNo": "EMP001",\n'
'      "email": "zhangsan@bank.com",\n'
'      "roles": ["admin"]\n'
'    }\n'
'  }\n'
'}', '成功响应')
resp_error(doc, [
    ['401', 'INVALID_CREDENTIALS', '用户名或密码错误'],
    ['400', 'ACCOUNT_DISABLED', '账号已被禁用'],
])

api_group(doc, '2.2 登出')
api_entry(doc, 'POST', '/api/auth/logout', '用户登出，使Token失效')
req_header(doc, [['Authorization', 'string', '是', 'Bearer {token}']])
resp_success(doc,
'{\n'
'  "code": 200,\n'
'  "message": "success",\n'
'  "data": null\n'
'}', '成功响应')

api_group(doc, '2.3 获取当前用户')
api_entry(doc, 'GET', '/api/auth/current-user', '获取当前登录用户信息')
req_header(doc, [['Authorization', 'string', '是', 'Bearer {token}']])
resp_success(doc,
'{\n'
'  "code": 200,\n'
'  "message": "success",\n'
'  "data": {\n'
'    "id": 1,\n'
'    "username": "EMP001",\n'
'    "name": "张三",\n'
'    "empNo": "EMP001",\n'
'    "email": "zhangsan@bank.com",\n'
'    "phone": "13800138000",\n'
'    "roles": ["admin"],\n'
'    "institutionIds": [],\n'
'    "status": 1\n'
'  }\n'
'}', '成功响应')

doc.add_page_break()
add_heading(doc, '3. 考核体系接口', 1)

api_group(doc, '3.1 体系列表')
api_entry(doc, 'GET', '/api/systems', '获取考核体系列表（分页+搜索）')
req_header(doc, [['Authorization', 'string', '是', 'Bearer {token}']])
req_param(doc, [
    ['name', 'query', 'string', '否', '体系名称模糊搜索'],
    ['status', 'query', 'int', '否', '状态：0-禁用，1-启用'],
    ['page', 'query', 'int', '否', '页码，默认1'],
    ['pageSize', 'query', 'int', '否', '每页条数，默认10'],
    ['sort', 'query', 'string', '否', '排序，默认createdAt,desc'],
])
resp_success(doc,
'{\n'
'  "code": 200,\n'
'  "message": "success",\n'
'  "data": {\n'
'    "records": [\n'
'      {\n'
'        "id": 1,\n'
'        "name": "2026年度分行绩效考核体系",\n'
'        "description": "适用于全行12家分行",\n'
'        "institutionCount": 12,\n'
'        "indicatorCount": 45,\n'
'        "needApproval": true,\n'
'        "status": 1,\n'
'        "createdBy": "张三",\n'
'        "createdAt": "2026-01-15T09:00:00.000+08:00",\n'
'        "updatedAt": "2026-01-15T09:00:00.000+08:00"\n'
'      }\n'
'    ],\n'
'    "total": 5,\n'
'    "page": 1,\n'
'    "pageSize": 10,\n'
'    "pages": 1\n'
'  }\n'
'}', '成功响应')

api_group(doc, '3.2 创建体系')
api_entry(doc, 'POST', '/api/systems', '创建考核体系（含Excel模版上传）')
req_header(doc, [['Authorization', 'string', '是', 'Bearer {token}']])
req_file(doc)
add_heading(doc, '其他表单字段', 4)
add_table(doc, ['字段', '类型', '必填', '说明'], [
    ['name', 'string', '是', '体系名称，最大100字符'],
    ['description', 'string', '否', '体系描述，最大500字符'],
    ['needApproval', 'boolean', '否', '是否需要审批，默认false'],
], col_widths=[3, 2, 1.5, 5])
resp_success(doc,
'{\n'
'  "code": 200,\n'
'  "message": "success",\n'
'  "data": {\n'
'    "id": 1,\n'
'    "name": "2026年度分行绩效考核体系",\n'
'    "institutionCount": 5,\n'
'    "indicatorCount": 24,\n'
'    "createdAt": "2026-03-27T10:30:00.000+08:00"\n'
'  }\n'
'}', '成功响应')
resp_error(doc, [
    ['400', 'INVALID_TEMPLATE', 'Excel模版格式不正确，缺少必要的Sheet或列'],
    ['400', 'TEMPLATE_PARSE_ERROR', 'Excel解析失败，请检查文件内容'],
    ['413', 'FILE_TOO_LARGE', '文件大小超过5MB'],
    ['415', 'UNSUPPORTED_FILE_TYPE', '仅支持.xlsx格式'],
])

api_group(doc, '3.3 体系详情')
api_entry(doc, 'GET', '/api/systems/{id}', '获取体系详情（含机构和指标列表）')
req_header(doc, [['Authorization', 'string', '是', 'Bearer {token}']])
req_param(doc, [['include', 'query', 'string', '否', '可选值：institutions, indicators，同时包含时逗号分隔']])
resp_success(doc,
'{\n'
'  "code": 200,\n'
'  "message": "success",\n'
'  "data": {\n'
'    "id": 1,\n'
'    "name": "2026年度分行绩效考核体系",\n'
'    "description": "适用于全行12家分行",\n'
'    "templateFileKey": "perf-templates/1/template_1743051000000.xlsx",\n'
'    "needApproval": true,\n'
'    "status": 1,\n'
'    "createdBy": "张三",\n'
'    "createdAt": "2026-01-15T09:00:00.000+08:00",\n'
'    "institutions": [\n'
'      {\n'
'        "id": 1,\n'
'        "orgName": "北京分行",\n'
'        "orgId": "ORG_BJ",\n'
'        "groupName": "北方区",\n'
'        "leaderName": "张三",\n'
'        "leaderEmpNo": "EMP001"\n'
'      }\n'
'    ],\n'
'    "indicators": [\n'
'      {\n'
'        "id": 1,\n'
'        "dimension": "业务发展",\n'
'        "category": "存款业务",\n'
'        "level1Name": "存款规模",\n'
'        "level2Name": "日均存款余额",\n'
'        "weight": 0.15,\n'
'        "unit": "亿元",\n'
'        "annualTarget": 100.00,\n'
'        "progressTarget": 8.33,\n'
'        "rowIndex": 3\n'
'      }\n'
'    ]\n'
'  }\n'
'}', '成功响应')

api_group(doc, '3.4 编辑体系')
api_entry(doc, 'PUT', '/api/systems/{id}', '编辑体系基本信息（模版不可修改）')
req_header(doc, [['Authorization', 'string', '是', 'Bearer {token}']])
req_body(doc, [
    ['name', 'string', '是', '体系名称'],
    ['description', 'string', '否', '体系描述'],
    ['needApproval', 'boolean', '否', '是否需要审批'],
    ['status', 'int', '否', '状态：0-禁用，1-启用'],
])
resp_success(doc,
'{\n'
'  "code": 200,\n'
'  "message": "success",\n'
'  "data": {\n'
'    "id": 1,\n'
'    "name": "2026年度分行绩效考核体系（修订版）",\n'
'    "updatedAt": "2026-03-27T11:00:00.000+08:00"\n'
'  }\n'
'}', '成功响应')
resp_error(doc, [
    ['403', 'SYSTEM_HAS_ACTIVE_MONITORING', '该体系下存在进行中的监测，无法编辑'],
])

api_group(doc, '3.5 删除体系')
api_entry(doc, 'DELETE', '/api/systems/{id}', '删除考核体系')
req_header(doc, [['Authorization', 'string', '是', 'Bearer {token}']])
resp_success(doc,
'{\n'
'  "code": 200,\n'
'  "message": "success",\n'
'  "data": null\n'
'}', '成功响应')
resp_error(doc, [
    ['403', 'SYSTEM_HAS_ACTIVE_MONITORING', '该体系下存在进行中的监测，无法删除'],
    ['404', 'SYSTEM_NOT_FOUND', '体系不存在'],
])

api_group(doc, '3.6 下载原始模版')
api_entry(doc, 'GET', '/api/systems/{id}/template/download', '下载该体系的原始Excel模版')
req_header(doc, [['Authorization', 'string', '是', 'Bearer {token}']])
doc.add_paragraph('响应：Content-Type: application/vnd.openxmlformats-officedocument.spreadsheetml.sheet，附带 Content-Disposition: attachment; filename="体系模版.xlsx"')

doc.add_page_break()
add_heading(doc, '4. 月度监测接口', 1)

api_group(doc, '4.1 监测列表')
api_entry(doc, 'GET', '/api/monitorings', '获取月度监测列表（分页+筛选）')
req_header(doc, [['Authorization', 'string', '是', 'Bearer {token}']])
req_param(doc, [
    ['systemId', 'query', 'long', '否', '体系ID'],
    ['year', 'query', 'int', '否', '年份'],
    ['month', 'query', 'int', '否', '月份'],
    ['status', 'query', 'string', '否', '状态筛选'],
    ['page', 'query', 'int', '否', '页码，默认1'],
    ['pageSize', 'query', 'int', '否', '每页条数，默认10'],
])
resp_success(doc,
'{\n'
'  "code": 200,\n'
'  "message": "success",\n'
'  "data": {\n'
'    "records": [\n'
'      {\n'
'        "id": 1,\n'
'        "systemId": 1,\n'
'        "systemName": "2026年度分行绩效考核体系",\n'
'        "year": 2026,\n'
'        "month": 3,\n'
'        "status": "COLLECTING",\n'
'        "deadline": "2026-03-27T18:00:00.000+08:00",\n'
'        "approvalRequired": true,\n'
'        "processPercent": 0,\n'
'        "processStatus": "idle",\n'
'        "createdBy": "管理员",\n'
'        "createdAt": "2026-03-27T08:00:00.000+08:00",\n'
'        "totalInstitutions": 12,\n'
'        "confirmedCount": 0,\n'
'        "pendingCount": 12\n'
'      }\n'
'    ],\n'
'    "total": 8,\n'
'    "page": 1,\n'
'    "pageSize": 10,\n'
'    "pages": 1\n'
'  }\n'
'}', '成功响应')

api_group(doc, '4.2 发起月度监测')
api_entry(doc, 'POST', '/api/monitorings', '发起新的月度监测')
req_header(doc, [['Authorization', 'string', '是', 'Bearer {token}']])
req_body(doc, [
    ['systemId', 'long', '是', '考核体系ID'],
    ['year', 'int', '是', '年份，如2026，范围当前年份-2~当前年份+1'],
    ['month', 'int', '是', '月份，1-12'],
    ['deadline', 'string', '是', '收数截止时间，ISO 8601格式'],
    ['approvalRequired', 'boolean', '否', '是否需要审批，不传则继承体系的needApproval'],
])
resp_success(doc,
'{\n'
'  "code": 200,\n'
'  "message": "success",\n'
'  "data": {\n'
'    "id": 2,\n'
'    "systemId": 1,\n'
'    "systemName": "2026年度分行绩效考核体系",\n'
'    "year": 2026,\n'
'    "month": 3,\n'
'    "status": "COLLECTING",\n'
'    "deadline": "2026-03-27T18:00:00.000+08:00"\n'
'  }\n'
'}', '成功响应')
resp_error(doc, [
    ['409', 'MONITORING_ALREADY_EXISTS', '该体系和月份的监测已存在，不可重复发起'],
    ['400', 'INVALID_DEADLINE', '截止时间不能早于当前时间'],
    ['404', 'SYSTEM_NOT_FOUND', '体系不存在'],
])

api_group(doc, '4.3 监测详情')
api_entry(doc, 'GET', '/api/monitorings/{id}', '获取月度监测详情')
req_header(doc, [['Authorization', 'string', '是', 'Bearer {token}']])
req_param(doc, [['include', 'query', 'string', '否', '可选值：institutions, indicators, tasks']])
resp_success(doc,
'{\n'
'  "code": 200,\n'
'  "message": "success",\n'
'  "data": {\n'
'    "id": 1,\n'
'    "systemId": 1,\n'
'    "systemName": "2026年度分行绩效考核体系",\n'
'    "year": 2026,\n'
'    "month": 3,\n'
'    "status": "CONFIRMING",\n'
'    "deadline": "2026-03-27T18:00:00.000+08:00",\n'
'    "approvalRequired": true,\n'
'    "processPercent": 100,\n'
'    "processStatus": "done",\n'
'    "processMsg": "已完成100个机构的报告生成",\n'
'    "totalInstitutions": 12,\n'
'    "confirmedCount": 5,\n'
'    "pendingCount": 7,\n'
'    "createdBy": "管理员",\n'
'    "createdAt": "2026-03-27T08:00:00.000+08:00",\n'
'    "updatedAt": "2026-03-27T10:30:00.000+08:00"\n'
'  }\n'
'}', '成功响应')

api_group(doc, '4.4 手动截止收数')
api_entry(doc, 'POST', '/api/monitorings/{id}/close', '手动截止收数，将状态从COLLECTING改为CLOSED')
req_header(doc, [['Authorization', 'string', '是', 'Bearer {token}']])
resp_success(doc,
'{\n'
'  "code": 200,\n'
'  "message": "success",\n'
'  "data": {\n'
'    "id": 1,\n'
'    "status": "CLOSED"\n'
'  }\n'
'}', '成功响应')
resp_error(doc, [
    ['400', 'INVALID_STATUS', '当前状态不允许截止操作'],
])

api_group(doc, '4.5 查询数据落库进度')
api_entry(doc, 'GET', '/api/monitorings/{id}/process/status', '查询数据落库进度（前端轮询调用）')
req_header(doc, [['Authorization', 'string', '是', 'Bearer {token}']])
resp_success(doc,
'{\n'
'  "code": 200,\n'
'  "message": "success",\n'
'  "data": {\n'
'    "monitoringId": 1,\n'
'    "processStatus": "processing",\n'
'    "processPercent": 65,\n'
'    "processMsg": "正在生成第65/100个机构的报告",\n'
'    "currentInstitution": "深圳分行",\n'
'    "totalInstitutions": 100,\n'
'    "processedCount": 65,\n'
'    "startTime": "2026-03-27T10:00:00.000+08:00",\n'
'    "estimatedTimeRemaining": "约30秒"\n'
'  }\n'
'}', '处理中响应')
resp_success(doc,
'{\n'
'  "code": 200,\n'
'  "message": "success",\n'
'  "data": {\n'
'    "monitoringId": 1,\n'
'    "processStatus": "done",\n'
'    "processPercent": 100,\n'
'    "processMsg": "已完成全部数据落库",\n'
'    "totalInstitutions": 100,\n'
'    "processedCount": 100,\n'
'    "startTime": "2026-03-27T10:00:00.000+08:00",\n'
'    "endTime": "2026-03-27T10:01:30.000+08:00"\n'
'  }\n'
'}', '处理完成响应')

api_group(doc, '4.6 机构负责人确认')
api_entry(doc, 'POST', '/api/monitorings/{id}/confirm', '机构负责人确认本机构数据')
req_header(doc, [['Authorization', 'string', '是', 'Bearer {token}']])
req_body(doc, [
    ['institutionId', 'long', '是', '机构ID'],
    ['confirmed', 'boolean', '是', '是否确认，true=确认，false=可填写备注驳回'],
    ['remark', 'string', '否', '确认备注，最大500字符'],
])
resp_success(doc,
'{\n'
'  "code": 200,\n'
'  "message": "success",\n'
'  "data": {\n'
'    "monitoringId": 1,\n'
'    "institutionId": 3,\n'
'    "confirmed": true,\n'
'    "confirmedAt": "2026-03-27T11:30:00.000+08:00"\n'
'  }\n'
'}', '成功响应')
resp_error(doc, [
    ['403', 'NOT_INSTITUTION_LEADER', '当前用户不是该机构的负责人'],
    ['400', 'INVALID_STATUS', '当前监测状态不允许确认操作'],
])

api_group(doc, '4.7 发布监测结果')
api_entry(doc, 'POST', '/api/monitorings/{id}/publish', '管理员手动发布监测结果')
req_header(doc, [['Authorization', 'string', '是', 'Bearer {token}']])
req_body(doc, [
    ['remark', 'string', '否', '发布备注，如"2026年3月绩效数据正式发布"'],
])
resp_success(doc,
'{\n'
'  "code": 200,\n'
'  "message": "success",\n'
'  "data": {\n'
'    "id": 1,\n'
'    "status": "PUBLISHED",\n'
'    "publishedAt": "2026-03-27T12:00:00.000+08:00",\n'
'    "publishedBy": "管理员"\n'
'  }\n'
'}', '成功响应')
resp_error(doc, [
    ['400', 'NOT_ALL_CONFIRMED', '仍有机构未确认：深圳分行、成都分行'],
    ['400', 'INVALID_STATUS', '当前状态不允许发布操作'],
])

api_group(doc, '4.8 获取确认进度')
api_entry(doc, 'GET', '/api/monitorings/{id}/confirm/progress', '获取各机构确认进度')
req_header(doc, [['Authorization', 'string', '是', 'Bearer {token}']])
resp_success(doc,
'{\n'
'  "code": 200,\n'
'  "message": "success",\n'
'  "data": {\n'
'    "monitoringId": 1,\n'
'    "totalInstitutions": 12,\n'
'    "confirmedCount": 10,\n'
'    "pendingCount": 2,\n'
'    "progress": 83,\n'
'    "institutions": [\n'
'      {\n'
'        "institutionId": 1,\n'
'        "orgName": "北京分行",\n'
'        "leaderName": "张三",\n'
'        "confirmed": true,\n'
'        "confirmedAt": "2026-03-27T11:00:00.000+08:00",\n'
'        "remark": null\n'
'      },\n'
'      {\n'
'        "institutionId": 2,\n'
'        "orgName": "上海分行",\n'
'        "leaderName": "李四",\n'
'        "confirmed": false,\n'
'        "confirmedAt": null,\n'
'        "remark": null\n'
'      }\n'
'    ]\n'
'  }\n'
'}', '成功响应')

doc.add_page_break()
add_heading(doc, '5. 数据收集接口', 1)

api_group(doc, '5.1 整体上传数据收集页')
api_entry(doc, 'POST', '/api/monitorings/{id}/collect/upload', '管理员整体上传数据收集页Excel，系统自动解析分发')
req_header(doc, [['Authorization', 'string', '是', 'Bearer {token}']])
req_file(doc)
add_heading(doc, '其他表单字段', 4)
add_table(doc, ['字段', '类型', '必填', '说明'], [
    ['remark', 'string', '否', '上传备注，如"3月第二周更新数据"'],
], col_widths=[3, 2, 1.5, 5])
resp_success(doc,
'{\n'
'  "code": 200,\n'
'  "message": "success",\n'
'  "data": {\n'
'    "uploadId": "upload_1743052000000",\n'
'    "monitoringId": 1,\n'
'    "fileKey": "perf-uploads/1/2026/3/upload_1743052000000.xlsx",\n'
'    "parsedCount": 60,\n'
'    "updatedTasks": 45,\n'
'    "newTasks": 15,\n'
'    "uploadedAt": "2026-03-27T14:00:00.000+08:00"\n'
'  }\n'
'}', '成功响应')
resp_error(doc, [
    ['400', 'INVALID_SHEET', '未找到数据收集页Sheet'],
    ['400', 'PARSE_ERROR', 'Excel解析失败，请检查数据格式'],
])

api_group(doc, '5.2 获取收数任务列表')
api_entry(doc, 'GET', '/api/monitorings/{id}/collect/tasks', '获取指定监测的收数任务列表')
req_header(doc, [['Authorization', 'string', '是', 'Bearer {token}']])
req_param(doc, [
    ['institutionId', 'query', 'long', '否', '按机构筛选'],
    ['status', 'query', 'string', '否', 'pending / submitted / approved / rejected'],
    ['collectorEmpNo', 'query', 'string', '否', '按收数人工号筛选'],
    ['page', 'query', 'int', '否', '页码'],
    ['pageSize', 'query', 'int', '否', '每页条数'],
])
resp_success(doc,
'{\n'
'  "code": 200,\n'
'  "message": "success",\n'
'  "data": {\n'
'    "records": [\n'
'      {\n'
'        "id": 101,\n'
'        "monitoringId": 1,\n'
'        "institutionId": 1,\n'
'        "orgName": "北京分行",\n'
'        "indicatorId": 3,\n'
'        "indicatorName": "日均存款余额",\n'
'        "dimension": "业务发展",\n'
'        "category": "存款业务",\n'
'        "unit": "亿元",\n'
'        "annualTarget": 100.00,\n'
'        "progressTarget": 8.33,\n'
'        "collectorName": "王五",\n'
'        "collectorEmpNo": "EMP005",\n'
'        "actualValue": 7.85,\n'
'        "status": "submitted",\n'
'        "submittedAt": "2026-03-27T10:30:00.000+08:00",\n'
'        "approvedBy": null,\n'
'        "approvedAt": null\n'
'      }\n'
'    ],\n'
'    "total": 120,\n'
'    "page": 1,\n'
'    "pageSize": 10\n'
'  }\n'
'}', '成功响应')

api_group(doc, '5.3 收数人提交任务数据')
api_entry(doc, 'PUT', '/api/monitorings/{id}/collect/tasks/{taskId}', '收数人提交或更新任务数据')
req_header(doc, [['Authorization', 'string', '是', 'Bearer {token}']])
req_body(doc, [
    ['actualValue', 'decimal', '是', '实际完成值'],
    ['remark', 'string', '否', '填写备注'],
])
resp_success(doc,
'{\n'
'  "code": 200,\n'
'  "message": "success",\n'
'  "data": {\n'
'    "taskId": 101,\n'
'    "status": "submitted",\n'
'    "actualValue": 7.85,\n'
'    "submittedAt": "2026-03-27T15:30:00.000+08:00"\n'
'  }\n'
'}', '成功响应')
resp_error(doc, [
    ['403', 'NOT_TASK_OWNER', '当前用户不是该任务的收数人'],
    ['400', 'INVALID_STATUS', '任务状态不允许修改（已审批或监测已截止）'],
    ['400', 'INVALID_VALUE', '实际值格式错误或超出合理范围'],
])

api_group(doc, '5.4 获取收数人的任务（按机构分组）')
api_entry(doc, 'GET', '/api/collect/tasks', '获取当前收数人的所有任务（按机构分组返回）')
req_header(doc, [['Authorization', 'string', '是', 'Bearer {token}']])
req_param(doc, [
    ['systemId', 'query', 'long', '否', '按体系筛选'],
    ['monitoringId', 'query', 'long', '否', '按监测筛选'],
    ['status', 'query', 'string', '否', 'pending / submitted'],
])
resp_success(doc,
'{\n'
'  "code": 200,\n'
'  "message": "success",\n'
'  "data": [\n'
'    {\n'
'      "monitoringId": 1,\n'
'      "systemName": "2026年度分行绩效考核体系",\n'
'      "month": "2026年3月",\n'
'      "deadline": "2026-03-27T18:00:00.000+08:00",\n'
'      "institution": {\n'
'        "id": 1,\n'
'        "orgName": "北京分行"\n'
'      },\n'
'      "tasks": [\n'
'        {\n'
'          "taskId": 101,\n'
'          "indicatorName": "日均存款余额",\n'
'          "unit": "亿元",\n'
'          "annualTarget": 100.00,\n'
'          "progressTarget": 8.33,\n'
'          "actualValue": null,\n'
'          "status": "pending"\n'
'        },\n'
'        {\n'
'          "taskId": 102,\n'
'          "indicatorName": "存款增长率",\n'
'          "unit": "%",\n'
'          "annualTarget": 15.00,\n'
'          "progressTarget": 1.25,\n'
'          "actualValue": 1.10,\n'
'          "status": "submitted",\n'
'          "submittedAt": "2026-03-27T10:30:00.000+08:00"\n'
'        }\n'
'      ]\n'
'    }\n'
'  ]\n'
'}', '成功响应')

api_group(doc, '5.5 审批收数任务')
api_entry(doc, 'PUT', '/api/monitorings/{id}/collect/tasks/{taskId}/approve', '管理员审批收数任务（需要审批的体系）')
req_header(doc, [['Authorization', 'string', '是', 'Bearer {token}']])
req_body(doc, [
    ['approved', 'boolean', '是', 'true=通过，false=驳回'],
    ['remark', 'string', '否', '审批备注，驳回时必填'],
])
resp_success(doc,
'{\n'
'  "code": 200,\n'
'  "message": "success",\n'
'  "data": {\n'
'    "taskId": 101,\n'
'    "status": "approved",\n'
'    "approvedBy": "管理员",\n'
'    "approvedAt": "2026-03-27T16:00:00.000+08:00"\n'
'  }\n'
'}', '审批通过响应')

doc.add_page_break()
add_heading(doc, '6. 结果展示接口', 1)

api_group(doc, '6.1 总览报表（所有机构对比）')
api_entry(doc, 'GET', '/api/reports/summary', '获取总览报表数据（所有机构汇总）')
req_header(doc, [['Authorization', 'string', '是', 'Bearer {token}']])
req_param(doc, [
    ['systemId', 'query', 'long', '是', '体系ID'],
    ['year', 'query', 'int', '是', '年份'],
    ['month', 'query', 'int', '是', '月份'],
    ['groupName', 'query', 'string', '否', '分组名称筛选'],
])
resp_success(doc,
'{\n'
'  "code": 200,\n'
'  "message": "success",\n'
'  "data": {\n'
'    "monitoringId": 1,\n'
'    "systemId": 1,\n'
'    "systemName": "2026年度分行绩效考核体系",\n'
'    "year": 2026,\n'
'    "month": 3,\n'
'    "groupName": null,\n'
'    "institutions": [\n'
'      {\n'
'        "institutionId": 1,\n'
'        "orgName": "北京分行",\n'
'        "groupName": "北方区",\n'
'        "totalScore": 92.3456,\n'
'        "dimensionScores": {\n'
'          "业务发展": 45.1234,\n'
'          "风险控制": 47.2222\n'
'        },\n'
'        "rank": 1\n'
'      },\n'
'      {\n'
'        "institutionId": 2,\n'
'        "orgName": "上海分行",\n'
'        "groupName": "华东区",\n'
'        "totalScore": 88.9012,\n'
'        "dimensionScores": {\n'
'          "业务发展": 42.3000,\n'
'          "风险控制": 46.6012\n'
'        },\n'
'        "rank": 2\n'
'      }\n'
'    ],\n'
'    "dimensions": ["业务发展", "风险控制", "服务质量", "合规管理"],\n'
'    "groups": ["北方区", "华东区", "华南区", "西南区"]\n'
'  }\n'
'}', '成功响应')

api_group(doc, '6.2 特定机构报表')
api_entry(doc, 'GET', '/api/reports/institution/{institutionId}', '获取特定机构的详细绩效数据')
req_header(doc, [['Authorization', 'string', '是', 'Bearer {token}']])
req_param(doc, [
    ['monitoringId', 'query', 'long', '否', '监测ID，不传则取最新已发布的'],
    ['year', 'query', 'int', '否', '年份，和monitoringId二选一'],
    ['month', 'query', 'int', '否', '月份'],
])
resp_success(doc,
'{\n'
'  "code": 200,\n'
'  "message": "success",\n'
'  "data": {\n'
'    "institutionId": 1,\n'
'    "orgName": "北京分行",\n'
'    "groupName": "北方区",\n'
'    "monitoringId": 1,\n'
'    "year": 2026,\n'
'    "month": 3,\n'
'    "totalScore": 92.3456,\n'
'    "indicators": [\n'
'      {\n'
'        "indicatorId": 1,\n'
'        "dimension": "业务发展",\n'
'        "category": "存款业务",\n'
'        "level1Name": "存款规模",\n'
'        "level2Name": "日均存款余额",\n'
'        "unit": "亿元",\n'
'        "weight": 0.15,\n'
'        "annualTarget": 100.00,\n'
'        "progressTarget": 8.33,\n'
'        "actualValue": 7.85,\n'
'        "annualCompletionRate": 0.0785,\n'
'        "progressCompletionRate": 0.9424,\n'
'        "score100": 94.24,\n'
'        "scoreWeighted": 14.136,\n'
'        "scoreCategory": 28.272,\n'
'        "scoreDimension": 45.1234,\n'
'        "fileKey": "perf-reports/1/2026/3/ORG_BJ/report_202603.xlsx"\n'
'      }\n'
'    ],\n'
'    "dimensionScores": {\n'
'      "业务发展": {"score": 45.1234, "weight": 0.6, "percentage": 75.21},\n'
'      "风险控制": {"score": 47.2222, "weight": 0.4, "percentage": 100.00}\n'
'    }\n'
'  }\n'
'}', '成功响应')

api_group(doc, '6.3 导出Excel报表')
api_entry(doc, 'GET', '/api/reports/export/excel', '根据筛选条件导出Excel绩效报表')
req_header(doc, [['Authorization', 'string', '是', 'Bearer {token}']])
req_param(doc, [
    ['systemId', 'query', 'long', '是', '体系ID'],
    ['year', 'query', 'int', '是', '年份'],
    ['month', 'query', 'int', '是', '月份'],
    ['institutionIds', 'query', 'string', '否', '机构ID，多个逗号分隔，不传则全部机构'],
])
doc.add_paragraph('响应：Content-Type: application/vnd.openxmlformats-officedocument.spreadsheetml.sheet，附带 Content-Disposition: attachment; filename="绩效报表_202603.xlsx"')

api_group(doc, '6.4 下载机构月度报告')
api_entry(doc, 'GET', '/api/reports/download/{monitoringId}/{institutionId}', '下载特定机构在特定监测中的月度报告Excel')
req_header(doc, [['Authorization', 'string', '是', 'Bearer {token}']])
doc.add_paragraph('响应：Content-Type: application/vnd.openxmlformats-officedocument.spreadsheetml.sheet，附带 Content-Disposition: attachment; filename="北京分行_2026年3月绩效报告.xlsx"')
resp_error(doc, [
    ['403', 'REPORT_NOT_PUBLISHED', '该监测结果尚未发布，仅发布后可下载'],
    ['403', 'NOT_AUTHORIZED', '机构负责人在发布前只能下载自己机构的数据'],
])

doc.add_page_break()
add_heading(doc, '7. 通知接口', 1)

api_group(doc, '7.1 通知列表')
api_entry(doc, 'GET', '/api/notifications', '获取当前用户的通知列表')
req_header(doc, [['Authorization', 'string', '是', 'Bearer {token}']])
req_param(doc, [
    ['read', 'query', 'boolean', '否', 'true=只看已读，false=只看未读，不传则全部'],
    ['page', 'query', 'int', '否', '页码，默认1'],
    ['pageSize', 'query', 'int', '否', '每页条数，默认10'],
])
resp_success(doc,
'{\n'
'  "code": 200,\n'
'  "message": "success",\n'
'  "data": {\n'
'    "records": [\n'
'      {\n'
'        "id": 1,\n'
'        "title": "【收数任务】您有新的数据待填写",\n'
'        "content": "体系：2026年度分行绩效考核体系，月份：2026年3月，请于03月27日18:00前完成。",\n'
'        "type": "site",\n'
'        "read": false,\n'
'        "actionUrl": "/collect",\n'
'        "actionParam": "{\"monitoringId\":1}",\n'
'        "createdAt": "2026-03-27T08:05:00.000+08:00"\n'
'      },\n'
'      {\n'
'        "id": 2,\n'
'        "title": "【截止提醒】收数即将截止",\n'
'        "content": "您的数据还未提交，距离截止还有30分钟，请尽快填写。",\n'
'        "type": "site",\n'
'        "read": true,\n'
'        "actionUrl": "/collect",\n'
'        "actionParam": "{\"monitoringId\":1}",\n'
'        "createdAt": "2026-03-27T17:30:00.000+08:00"\n'
'      }\n'
'    ],\n'
'    "total": 15,\n'
'    "page": 1,\n'
'    "pageSize": 10\n'
'  }\n'
'}', '成功响应')

api_group(doc, '7.2 标记单条已读')
api_entry(doc, 'PUT', '/api/notifications/{id}/read', '标记指定通知为已读')
req_header(doc, [['Authorization', 'string', '是', 'Bearer {token}']])
resp_success(doc,
'{\n'
'  "code": 200,\n'
'  "message": "success",\n'
'  "data": {\n'
'    "id": 1,\n'
'    "read": true,\n'
'    "readAt": "2026-03-27T14:00:00.000+08:00"\n'
'  }\n'
'}', '成功响应')

api_group(doc, '7.3 全部标记已读')
api_entry(doc, 'PUT', '/api/notifications/read-all', '将当前用户所有未读通知全部标记为已读')
req_header(doc, [['Authorization', 'string', '是', 'Bearer {token}']])
resp_success(doc,
'{\n'
'  "code": 200,\n'
'  "message": "success",\n'
'  "data": {\n'
'    "totalRead": 14\n'
'  }\n'
'}', '成功响应')

api_group(doc, '7.4 未读数量')
api_entry(doc, 'GET', '/api/notifications/unread-count', '获取当前用户未读通知数量（用于Badge显示）')
req_header(doc, [['Authorization', 'string', '是', 'Bearer {token}']])
resp_success(doc,
'{\n'
'  "code": 200,\n'
'  "message": "success",\n'
'  "data": {\n'
'    "unreadCount": 1\n'
'  }\n'
'}', '成功响应')

doc.add_page_break()
add_heading(doc, '8. 仪表盘接口', 1)

api_group(doc, '8.1 仪表盘统计数据')
api_entry(doc, 'GET', '/api/dashboard/stats', '获取仪表盘统计数据（不同角色返回不同字段）')
req_header(doc, [['Authorization', 'string', '是', 'Bearer {token}']])
resp_success(doc,
'{\n'
'  "code": 200,\n'
'  "message": "success",\n'
'  "data": {\n'
'    "activeMonitoringCount": 3,\n'
'    "pendingConfirmCount": 1,\n'
'    "publishedCount": 12,\n'
'    "myPendingTaskCount": 5,\n'
'    "recentMonitoring": [\n'
'      {\n'
'        "id": 3,\n'
'        "systemName": "2026年度分行绩效考核体系",\n'
'        "year": 2026,\n'
'        "month": 3,\n'
'        "status": "COLLECTING",\n'
'        "deadline": "2026-03-27T18:00:00.000+08:00"\n'
'      }\n'
'    ]\n'
'  }\n'
'}', '管理员响应示例')
resp_success(doc,
'{\n'
'  "code": 200,\n'
'  "message": "success",\n'
'  "data": {\n'
'    "myInstitutionName": "北京分行",\n'
'    "myPendingTaskCount": 5,\n'
'    "pendingConfirmCount": 1,\n'
'    "recentMonitoring": [\n'
'      {\n'
'        "id": 3,\n'
'        "systemName": "2026年度分行绩效考核体系",\n'
'        "year": 2026,\n'
'        "month": 3,\n'
'        "status": "COLLECTING",\n'
'        "deadline": "2026-03-27T18:00:00.000+08:00"\n'
'      }\n'
'    ]\n'
'  }\n'
'}', '收数人/机构负责人响应示例')

api_group(doc, '8.2 最近动态')
api_entry(doc, 'GET', '/api/dashboard/recent', '获取最近操作动态（时间线）')
req_header(doc, [['Authorization', 'string', '是', 'Bearer {token}']])
req_param(doc, [['limit', 'query', 'int', '否', '返回条数，默认10，最大50']])
resp_success(doc,
'{\n'
'  "code": 200,\n'
'  "message": "success",\n'
'  "data": [\n'
'    {\n'
'      "id": 1,\n'
'      "type": "monitoring_published",\n'
'      "content": "管理员发布了2026年2月的绩效结果",\n'
'      "operator": "管理员",\n'
'      "operatorAvatar": null,\n'
'      "actionTime": "2026-03-01T12:00:00.000+08:00"\n'
'    },\n'
'    {\n'
'      "id": 2,\n'
'      "type": "task_submitted",\n'
'      "content": "王五提交了北京分行3月的存款数据",\n'
'      "operator": "王五",\n'
'      "operatorAvatar": null,\n'
'      "actionTime": "2026-03-27T10:30:00.000+08:00"\n'
'    }\n'
'  ]\n'
'}', '成功响应')

doc.add_page_break()
add_heading(doc, '附录：错误码完整对照表', 1)
add_table(doc, ['错误码', 'HTTP状态码', '说明', '处理建议'], [
    ['INVALID_CREDENTIALS', '401', '用户名或密码错误', '提示用户检查用户名密码'],
    ['ACCOUNT_DISABLED', '401', '账号已被禁用', '联系管理员启用账号'],
    ['TOKEN_EXPIRED', '401', 'Token已过期', '重新登录'],
    ['TOKEN_INVALID', '401', 'Token无效', '重新登录'],
    ['ACCESS_DENIED', '403', '无权限访问该资源', '检查用户角色和权限'],
    ['NOT_TASK_OWNER', '403', '不是任务的收数人', '当前用户无权操作该任务'],
    ['NOT_INSTITUTION_LEADER', '403', '不是该机构的负责人', '当前用户无权确认该机构数据'],
    ['REPORT_NOT_PUBLISHED', '403', '报告尚未发布', '发布后再操作'],
    ['SYSTEM_NOT_FOUND', '404', '体系不存在', '检查体系ID'],
    ['MONITORING_NOT_FOUND', '404', '监测不存在', '检查监测ID'],
    ['TASK_NOT_FOUND', '404', '任务不存在', '检查任务ID'],
    ['INVALID_TEMPLATE', '400', 'Excel模版格式不正确', '按模版要求修改后重新上传'],
    ['TEMPLATE_PARSE_ERROR', '400', 'Excel解析失败', '检查文件内容是否完整'],
    ['INVALID_DEADLINE', '400', '截止时间设置不合理', '截止时间需晚于当前时间'],
    ['INVALID_STATUS', '400', '当前状态不允许该操作', '检查监测当前状态'],
    ['MONITORING_ALREADY_EXISTS', '409', '同一月份监测已存在', '不可重复发起'],
    ['SYSTEM_HAS_ACTIVE_MONITORING', '409', '体系下存在进行中的监测', '先完成或取消进行中的监测'],
    ['NOT_ALL_CONFIRMED', '400', '仍有机构未确认', '确认所有机构后再发布'],
    ['FILE_TOO_LARGE', '413', '上传文件超过5MB', '压缩文件或分批上传'],
    ['UNSUPPORTED_FILE_TYPE', '415', '不支持的文件格式', '仅支持.xlsx格式'],
], col_widths=[3.5, 2, 4, 3])

doc.save('/Users/zhaoyu/Desktop/机构绩效管理系统_API详细设计.docx')
print('API detailed design doc generated successfully')
