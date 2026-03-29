#!/usr/bin/env python3
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

def set_cell_bg(cell, color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)

def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        set_cell_bg(cell, '2E4057')
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(10)
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        bg = 'F8F9FA' if ri % 2 == 0 else 'FFFFFF'
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = str(val)
            set_cell_bg(cell, bg)
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in para.runs:
                    run.font.size = Pt(9)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    return table

def add_image(doc, img_path, width=None, height=None):
    from docx.shared import Inches
    if width:
        doc.add_picture(img_path, width=width)
    elif height:
        doc.add_picture(img_path, height=height)
    else:
        doc.add_picture(img_path)
    last_para = doc.paragraphs[-1]
    last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

def section(doc, title):
    add_heading(doc, title, 1)
    doc.add_page_break()

doc = Document()
section = doc.sections[0]
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)

title = doc.add_heading('机构绩效管理系统\n页面设计文档', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph(f'版本：V1.0.0\n日期：{datetime.date.today().strftime("%Y-%m-%d")}\n状态：初稿').alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()
add_heading(doc, '目录', 1)
doc.add_paragraph(
    '1. 设计规范\n'
    '2. 页面清单与路由\n'
    '3. 登录页 /login\n'
    '4. 首页仪表盘 /dashboard\n'
    '5. 考核体系列表页 /system\n'
    '6. 发起监测弹窗\n'
    '7. 监测详情页 /monitoring/:id\n'
    '8. 数据填写页（收数人）/collect\n'
    '9. 绩效报表页 /report\n'
    '10. 通知中心页 /notifications\n'
    '11. 通用组件规范'
)

doc.add_page_break()
add_heading(doc, '1. 设计规范', 1)
add_heading(doc, '1.1 设计风格', 2)
add_table(doc, ['项目', '规范'], [
    ['设计语言', '企业级B端产品风格，简洁专业'],
    ['组件库', 'Ant Design 6.x'],
    ['色彩体系', '主色 #2E4057（深蓝），辅色 #048A81（青绿），警示色 #E07A5F（珊瑚红）'],
    ['字体', 'PingFang SC / Microsoft YaHei'],
    ['图标', 'Ant Design Icons'],
    ['图表', 'ECharts 6.x'],
    ['布局', '侧边导航 Layout + 内容区 Content'],
], col_widths=[3, 8])

add_heading(doc, '1.2 布局结构', 2)
add_table(doc, ['区域', '高度/宽度', '说明'], [
    ['顶部导航栏 Header', '64px', 'Logo、系统名称、当前用户信息、通知Badge、下拉菜单'],
    ['侧边导航栏 Sider', '240px（可折叠至64px）', '根据角色动态渲染菜单项'],
    ['内容区 Content', '自适应', '白色背景，内边距 24px'],
    ['页面标题区', '约56px', '当前页面名称 + 操作按钮'],
], col_widths=[3, 2.5, 6])

add_heading(doc, '1.3 状态色彩规范', 2)
add_table(doc, ['状态', '色彩', '用途'], [
    ['待发起/PENDING', '#999999 灰色', 'Tag标签'],
    ['收数中/COLLECTING', '#1890FF 蓝色', 'Tag标签'],
    ['已截止/CLOSED', '#FA8C16 橙色', 'Tag标签'],
    ['数据落库中/PROCESSING', '#722ED1 紫色', 'Tag标签'],
    ['待确认/CONFIRMING', '#FA8C16 橙色', 'Tag标签'],
    ['已发布/PUBLISHED', '#52C41A 绿色', 'Tag标签'],
], col_widths=[3, 3, 5])

doc.add_page_break()
add_heading(doc, '2. 页面清单与路由', 1)
add_table(doc, ['页面', '路由', '角色', '优先级'], [
    ['登录页', '/login', '公开', 'P0'],
    ['首页仪表盘', '/dashboard', '管理员、机构负责人', 'P0'],
    ['考核体系列表', '/system', '管理员', 'P0'],
    ['创建体系', '/system/create', '管理员', 'P0'],
    ['月度监测列表', '/monitoring', '管理员', 'P0'],
    ['监测详情', '/monitoring/:id', '管理员', 'P0'],
    ['数据填写（收数人）', '/collect', '收数人', 'P0'],
    ['绩效报表', '/report', '管理员、机构负责人（发布后）', 'P0'],
    ['总览对比', '/report/overview', '管理员、机构负责人（发布后）', 'P0'],
    ['通知中心', '/notifications', '所有登录用户', 'P1'],
], col_widths=[4, 4, 3, 1.5])

doc.add_page_break()
add_heading(doc, '3. 登录页 /login', 1)

add_heading(doc, '3.1 页面说明', 2)
doc.add_paragraph('路由：/login | 权限：公开 | 功能：用户登录系统，获取JWT Token。')

add_heading(doc, '3.2 UI设计图', 2)
doc.add_picture('/Users/zhaoyu/Desktop/UI_Mockups/01_login.svg.png', width=Cm(16))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('图3-1 登录页 UI设计（1200×800px）').alignment = WD_ALIGN_PARAGRAPH.CENTER

add_heading(doc, '3.3 布局说明', 2)
add_table(doc, ['区域', '元素', '说明'], [
    ['左侧品牌区（深蓝底色）', 'Logo + 系统名称 + Slogan', '固定展示'],
    ['右侧表单区', '用户名输入框、密码输入框、登录按钮', '垂直排列'],
    ['用户名输入框', 'Input', 'placeholder="请输入工号"，必填'],
    ['密码输入框', 'Input.Password', 'placeholder="请输入密码"，必填，支持回车提交'],
    ['登录按钮', 'Button（主按钮）', '点击后显示loading，成功后跳转/dashboard'],
], col_widths=[3, 4, 5])

doc.add_page_break()
add_heading(doc, '4. 首页仪表盘 /dashboard', 1)

add_heading(doc, '4.1 页面说明', 2)
doc.add_paragraph('路由：/dashboard | 权限：管理员、机构负责人 | 功能：展示关键数据统计、快捷入口和最近动态。')

add_heading(doc, '4.2 UI设计图', 2)
doc.add_picture('/Users/zhaoyu/Desktop/UI_Mockups/02_dashboard.svg.png', width=Cm(16))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('图4-1 首页仪表盘 UI设计（1200×800px）').alignment = WD_ALIGN_PARAGRAPH.CENTER

add_heading(doc, '4.3 布局说明', 2)
add_table(doc, ['区域', '元素', '说明'], [
    ['统计卡片区', '4个统计卡片', '进行中监测数（蓝色）/ 待确认数（橙色）/ 已发布数（绿色）/ 我的待办数（紫色）'],
    ['快捷入口区', '4个快捷按钮', '发起月度监测 / 填写绩效数据 / 查看绩效报表 / 通知中心'],
    ['最近动态区', '时间线列表', '显示最近操作记录，含时间、操作人、内容'],
], col_widths=[3, 4, 5])

doc.add_page_break()
add_heading(doc, '5. 考核体系列表页 /system', 1)

add_heading(doc, '5.1 页面说明', 2)
doc.add_paragraph('路由：/system | 权限：管理员 | 功能：展示所有考核体系列表，支持搜索、分页、创建、编辑、删除操作。')

add_heading(doc, '5.2 UI设计图', 2)
doc.add_picture('/Users/zhaoyu/Desktop/UI_Mockups/03_system_list.svg.png', width=Cm(16))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('图5-1 考核体系列表页 UI设计（1200×800px）').alignment = WD_ALIGN_PARAGRAPH.CENTER

add_heading(doc, '5.3 布局说明', 2)
add_table(doc, ['区域', '元素', '说明'], [
    ['搜索筛选区', '体系名称搜索框 + 状态下拉筛选 + "创建体系"按钮', '搜索实时过滤'],
    ['数据表格', '体系名称/描述/机构数/指标数/状态/创建时间/操作', '分页展示，每页10条'],
    ['操作列', '编辑 / 查看 / 删除', '删除需二次确认'],
    ['状态标签', '启用（绿色）/ 禁用（灰色）', 'Tag组件'],
], col_widths=[3, 4, 5])

doc.add_page_break()
add_heading(doc, '6. 发起监测弹窗', 1)

add_heading(doc, '6.1 弹窗说明', 2)
doc.add_paragraph('触发位置：月度监测列表页 /monitoring，点击"发起监测"按钮弹出。')

add_heading(doc, '6.2 UI设计图', 2)
doc.add_picture('/Users/zhaoyu/Desktop/UI_Mockups/04_monitoring_create.svg.png', width=Cm(16))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('图6-1 发起监测弹窗 UI设计（1200×800px）').alignment = WD_ALIGN_PARAGRAPH.CENTER

add_heading(doc, '6.3 布局说明', 2)
add_table(doc, ['字段', '类型', '说明'], [
    ['选择考核体系', 'Select', '必选，下拉选择已有体系'],
    ['年份', 'InputNumber', '必选，默认当前年份'],
    ['月份', 'InputNumber', '必选，默认当前月份'],
    ['收数截止时间', 'DatePicker', '必选，支持时分选择'],
    ['是否需要审批', 'Switch', '默认继承体系的needApproval'],
    ['取消按钮', 'Default Button', '关闭弹窗'],
    ['确认发起按钮', 'Primary Button', '校验后提交'],
], col_widths=[3, 2.5, 6])

doc.add_page_break()
add_heading(doc, '7. 监测详情页 /monitoring/:id', 1)

add_heading(doc, '7.1 页面说明', 2)
doc.add_paragraph('路由：/monitoring/:id | 权限：管理员 | 功能：展示特定月度监测的完整信息，包括步骤条、信息卡片、数据收集表格。')

add_heading(doc, '7.2 UI设计图', 2)
doc.add_picture('/Users/zhaoyu/Desktop/UI_Mockups/05_monitoring_detail.svg.png', width=Cm(16))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('图7-1 监测详情页 UI设计（1200×800px）').alignment = WD_ALIGN_PARAGRAPH.CENTER

add_heading(doc, '7.3 布局说明', 2)
add_table(doc, ['区域', '元素', '说明'], [
    ['步骤条', '5个步骤横向排列', '发起→收数中→数据落库→待确认→已发布，当前步骤高亮'],
    ['信息卡片', '6个指标卡片', '体系名称/月份/状态/截止时间/待填写数/已提交数'],
    ['Tab区', '数据收集 / 落库进度 / 确认情况 / 关联文件', 'Tab切换'],
    ['数据收集表格', '机构/指标/收数人/实际值/状态/提交时间', '含"整体上传"按钮'],
    ['状态标签', '已提交（绿色）/ 待填写（灰色）', 'Tag组件'],
], col_widths=[3, 4, 5])

doc.add_page_break()
add_heading(doc, '8. 数据填写页（收数人） /collect', 1)

add_heading(doc, '8.1 页面说明', 2)
doc.add_paragraph('路由：/collect | 权限：收数人 | 功能：收数人查看自己被分配的指标任务，填写实际值后一次性提交。')

add_heading(doc, '8.2 UI设计图', 2)
doc.add_picture('/Users/zhaoyu/Desktop/UI_Mockups/06_collect.svg.png', width=Cm(16))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('图8-1 数据填写页 UI设计（1200×800px）').alignment = WD_ALIGN_PARAGRAPH.CENTER

add_heading(doc, '8.3 布局说明', 2)
add_table(doc, ['区域', '元素', '说明'], [
    ['机构分组卡片', '每个机构一个卡片', '包含机构名/体系名/月份/截止时间'],
    ['指标列表', '指标名称/单位/进度目标/全年目标/实际值输入框', '按机构分组展示'],
    ['操作按钮', '下载Excel / 上传Excel / 提交数据', '填写完成后才可提交'],
    ['状态标识', '待填写（灰色标签）/ 已提交（绿色标签）', 'Tag组件'],
], col_widths=[3, 4, 5])

doc.add_page_break()
add_heading(doc, '9. 绩效报表页 /report', 1)

add_heading(doc, '9.1 页面说明', 2)
doc.add_paragraph('路由：/report | 权限：管理员、机构负责人（发布后） | 功能：展示绩效数据，支持筛选、图表可视化和Excel导出。')

add_heading(doc, '9.2 UI设计图', 2)
doc.add_picture('/Users/zhaoyu/Desktop/UI_Mockups/07_report.svg.png', width=Cm(16))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('图9-1 绩效报表页 UI设计（1200×800px）').alignment = WD_ALIGN_PARAGRAPH.CENTER

add_heading(doc, '9.3 布局说明', 2)
add_table(doc, ['区域', '元素', '说明'], [
    ['筛选器栏', '体系/年份/月份/机构/分组 筛选器', '5个下拉选择器 + 重置按钮'],
    ['Tab区', '数据表格 / 可视化图表 / 总览对比', '默认展示可视化图表'],
    ['机构得分排名柱状图', '横向柱状图', '各机构总得分降序排列'],
    ['维度得分雷达图', '雷达图', '多机构多维度得分叠加对比'],
    ['指标完成率进度条', '进度条列表', '各指标实际完成百分比'],
], col_widths=[3, 4, 5])

doc.add_page_break()
add_heading(doc, '10. 通知中心页 /notifications', 1)

add_heading(doc, '10.1 页面说明', 2)
doc.add_paragraph('路由：/notifications | 权限：所有登录用户 | 功能：展示当前用户收到的站内通知，支持标记已读。')

add_heading(doc, '10.2 UI设计图', 2)
doc.add_picture('/Users/zhaoyu/Desktop/UI_Mockups/08_notifications.svg.png', width=Cm(16))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('图10-1 通知中心页 UI设计（1200×800px）').alignment = WD_ALIGN_PARAGRAPH.CENTER

add_heading(doc, '10.3 布局说明', 2)
add_table(doc, ['区域', '元素', '说明'], [
    ['Tab区', '全部 / 未读', '未读Tab显示Badge数量'],
    ['通知卡片', '图标/标题/内容/时间/操作', '未读通知左侧有蓝色竖条'],
    ['操作', '标为已读 / 查看详情', '点击跳转到对应业务页面'],
    ['全部标为已读', 'Text按钮', '一键已读所有通知'],
    ['分页器', '底部分页', '每页10条'],
], col_widths=[3, 4, 5])

doc.add_page_break()
add_heading(doc, '11. 通用组件规范', 1)

add_heading(doc, '11.1 状态标签', 2)
add_table(doc, ['监测状态', '色彩', '适用场景'], [
    ['PENDING（待发起）', '#999999 灰色', '发起监测前'],
    ['COLLECTING（收数中）', '#1890FF 蓝色', '数据收集阶段'],
    ['CLOSED（已截止）', '#FA8C16 橙色', '收数截止后'],
    ['PROCESSING（数据落库中）', '#722ED1 紫色', '异步处理中'],
    ['CONFIRMING（待确认）', '#FA8C16 橙色', '负责人确认阶段'],
    ['PUBLISHED（已发布）', '#52C41A 绿色', '结果发布后'],
], col_widths=[3, 3, 6])

add_heading(doc, '11.2 确认对话框', 2)
doc.add_paragraph('所有危险操作（删除、发布、截止）需弹出 Modal.confirm 确认框，包含操作描述 + 取消/确认按钮。')

add_heading(doc, '11.3 加载状态', 2)
add_table(doc, ['场景', '处理方式'], [
    ['页面初始加载', 'Spin加载动画遮罩整个内容区'],
    ['按钮操作中', 'Button设置loading=true，显示spinner'],
    ['文件上传中', 'Upload组件内置Progress进度条'],
    ['数据落库进度', 'Progress组件实时展示百分比'],
], col_widths=[3, 8])

add_heading(doc, '11.4 权限控制', 2)
for r in [
    '未登录访问任何页面 -> 跳转 /login',
    '收数人访问 /system 或 /monitoring -> 跳转 /dashboard 并提示"无权限"',
    '机构负责人发布前访问 /report -> 仅展示其关联机构数据',
    '机构负责人发布后访问 /report -> 可查看所有机构数据',
]:
    doc.add_paragraph(r, style='List Bullet')

doc.add_page_break()
add_heading(doc, '附录：UI设计图清单', 1)
add_table(doc, ['序号', '页面', '文件名'], [
    ['图3-1', '登录页 /login', '01_login.svg.png'],
    ['图4-1', '首页仪表盘 /dashboard', '02_dashboard.svg.png'],
    ['图5-1', '考核体系列表页 /system', '03_system_list.svg.png'],
    ['图6-1', '发起监测弹窗', '04_monitoring_create.svg.png'],
    ['图7-1', '监测详情页 /monitoring/:id', '05_monitoring_detail.svg.png'],
    ['图8-1', '数据填写页 /collect', '06_collect.svg.png'],
    ['图9-1', '绩效报表页 /report', '07_report.svg.png'],
    ['图10-1', '通知中心页 /notifications', '08_notifications.svg.png'],
], col_widths=[1.5, 4, 5])

doc.save('/Users/zhaoyu/Desktop/机构绩效管理系统_页面设计文档.docx')
print('Page design doc with UI mockups generated successfully')
