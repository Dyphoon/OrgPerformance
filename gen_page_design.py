#!/usr/bin/env python3
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

def set_cell_bg(cell, color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)

def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)

def add_para(doc, text, bold=False, italic=False):
    p = doc.add_paragraph(text)
    if bold or italic:
        for run in p.runs:
            run.bold = bold
            run.italic = italic
    return p

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        set_cell_bg(cell, '2E4057')
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(10)
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        bg = 'F8F9FA' if ri % 2 == 0 else 'FFFFFF'
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = str(val)
            set_cell_bg(cell, bg)
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in para.runs:
                    run.font.size = Pt(9)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    return table

def add_page_header(doc, page_num, page_name, path, role, desc):
    add_heading(doc, f'{page_num}. {page_name}', 2)
    p = doc.add_paragraph()
    p.add_run('路由：').bold = True
    p.add_run(path)
    p2 = doc.add_paragraph()
    p2.add_run('权限：').bold = True
    p2.add_run(role)
    p3 = doc.add_paragraph()
    p3.add_run('功能描述：').bold = True
    p3.add_run(desc)

doc = Document()
section = doc.sections[0]
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)

title = doc.add_heading('机构绩效管理系统\n页面设计文档', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph(f'版本：V1.0.0\n日期：{datetime.date.today().strftime("%Y-%m-%d")}\n状态：初稿').alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_page_break()

add_heading(doc, '目录', 1)
doc.add_paragraph(
    '1. 设计规范\n'
    '2. 页面清单与路由\n'
    '3. 登录页 /login\n'
    '4. 首页仪表盘 /dashboard\n'
    '5. 考核体系列表页 /system\n'
    '6. 创建/编辑体系页 /system/create\n'
    '7. 月度监测列表页 /monitoring\n'
    '8. 监测详情页 /monitoring/:id\n'
    '9. 数据填写页 /collect\n'
    '10. 绩效报表页 /report\n'
    '11. 总览对比页 /report/overview\n'
    '12. 通知中心页 /notifications\n'
    '13. 通用组件规范\n'
    '14. 响应式与无障碍'
)

doc.add_page_break()
add_heading(doc, '1. 设计规范', 1)

add_heading(doc, '1.1 设计风格', 2)
add_table(doc, ['项目', '规范'], [
    ['设计语言', '企业级 B端产品风格，简洁专业'],
    ['组件库', 'Ant Design 6.x'],
    ['色彩体系', '主色 #2E4057（深蓝），辅色 #048A81（青绿），警示色 #E07A5F（珊瑚红）'],
    ['字体', '系统默认字体，标题使用 PingFang SC / Microsoft YaHei'],
    ['图标', 'Ant Design Icons'],
    ['图表', 'ECharts 6.x'],
    ['布局', '侧边导航 Layout + 内容区 Content'],
    ['响应式', '最小宽度 1200px，针对 B端管理场景以桌面端为主'],
], col_widths=[3, 8])

add_heading(doc, '1.2 布局结构', 2)
add_table(doc, ['区域', '高度/宽度', '说明'], [
    ['顶部导航栏 Header', '64px', 'Logo、系统名称、当前用户信息、下拉菜单（通知、个人设置、退出）'],
    ['侧边导航栏 Sider', '240px（可折叠至64px）', '根据角色动态渲染菜单项，支持折叠展开'],
    ['内容区 Content', '自适应', '白色背景，内边距 24px，flex 布局'],
    ['页面标题区', '约56px', '当前页面名称面包屑 + 操作按钮'],
], col_widths=[3, 2.5, 6])

add_heading(doc, '1.3 状态色彩规范', 2)
add_table(doc, ['状态', '色彩', '用途'], [
    ['待发起/PENDING', '#999999 灰色', '文字标签'],
    ['收数中/COLLECTING', '#1890FF 蓝色', '文字标签'],
    ['已截止/CLOSED', '#FA8C16 橙色', '文字标签'],
    ['数据落库中/PROCESSING', '#722ED1 紫色', '文字标签'],
    ['待确认/CONFIRMING', '#FA8C16 橙色', '文字标签'],
    ['已发布/PUBLISHED', '#52C41A 绿色', '文字标签'],
    ['成功/SUCCESS', '#52C41A', '表单提交成功提示'],
    ['警告/WARNING', '#FA8C16', '表单校验警告'],
    ['错误/ERROR', '#F5222D', '表单错误、接口失败'],
    ['信息/INFO', '#1890FF', '通知提示'],
], col_widths=[3, 3, 5])

add_heading(doc, '1.4 表格规范', 2)
items = [
    '列表页统一使用 Ant Design Table 组件，支持排序、筛选、分页',
    '分页器默认每页 10 条，支持 10/20/50/100 条切换',
    '操作列固定在最后一列，宽度 120px',
    '状态列使用 Tag 组件，颜色按状态色彩规范',
    '日期列格式：YYYY-MM-DD HH:mm'],
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()
add_heading(doc, '2. 页面清单与路由', 1)

add_table(doc, ['页面', '路由', '角色', '优先级'], [
    ['登录页', '/login', '公开', 'P0'],
    ['首页仪表盘', '/dashboard', '管理员、机构负责人', 'P0'],
    ['考核体系列表', '/system', '管理员', 'P0'],
    ['创建体系', '/system/create', '管理员', 'P0'],
    ['编辑体系', '/system/:id/edit', '管理员', 'P0'],
    ['月度监测列表', '/monitoring', '管理员', 'P0'],
    ['监测详情', '/monitoring/:id', '管理员', 'P0'],
    ['数据填写（收数人）', '/collect', '收数人', 'P0'],
    ['绩效报表', '/report', '管理员、机构负责人（发布后）', 'P0'],
    ['总览对比', '/report/overview', '管理员、机构负责人（发布后）', 'P0'],
    ['通知中心', '/notifications', '所有登录用户', 'P1'],
    ['体系详情', '/system/:id', '管理员', 'P1'],
], col_widths=[4, 4, 3, 1.5])

doc.add_page_break()
add_heading(doc, '3. 登录页 /login', 1)

add_page_header(doc, '3', '登录页', '/login', '公开', '用户输入用户名密码登录系统，获取JWT Token并跳转首页。')

add_heading(doc, '3.1 布局', 2)
doc.add_paragraph('单页设计，居中卡片式布局。左侧为系统介绍区（含Logo、系统名称、Slogan），右侧为登录表单区。')

add_heading(doc, '3.2 页面元素', 2)
add_table(doc, ['元素', '类型', '说明'], [
    ['Logo + 系统名称', '图片+文字', '固定在左侧卡片顶部'],
    ['系统Slogan', '文本', '"机构绩效管理，数字化赋能分行考核"'],
    ['用户名输入框', 'Input', 'placeholder="请输入工号"，必填，最大长度50'],
    ['密码输入框', 'Input.Password', 'placeholder="请输入密码"，必填，支持回车提交'],
    ['登录按钮', 'Button', '主按钮，文本"登录"，点击触发登录请求'],
    ['错误提示', 'Alert', '登录失败时显示错误信息'],
], col_widths=[3, 2.5, 6])

add_heading(doc, '3.3 交互逻辑', 2)
steps = [
    '用户输入用户名和密码，点击登录或按回车',
    '前端校验：用户名为空或密码为空时，显示"请输入用户名/密码"提示，不发请求',
    '校验通过后，按钮显示 loading 状态（防重复提交）',
    '调用 POST /api/auth/login，成功后：存储 Token 到 localStorage，跳转 /dashboard',
    '失败后：按钮恢复，显示错误 Alert（用户名或密码错误）',
    'JWT Token 有效期 24 小时，过期后自动跳转登录页',
]
for s in steps:
    doc.add_paragraph(s, style='List Bullet')

add_heading(doc, '3.4 关联API', 2)
add_table(doc, ['接口', '方法', '说明'], [
    ['POST /api/auth/login', 'POST', '登录认证'],
], col_widths=[4, 1.5, 6])

doc.add_page_break()
add_heading(doc, '4. 首页仪表盘 /dashboard', 1)

add_page_header(doc, '4', '首页仪表盘', '/dashboard', '管理员、机构负责人', '展示关键数据统计、快捷入口和最近动态。不同角色看到的内容有所差异。')

add_heading(doc, '4.1 布局', 2)
doc.add_paragraph('顶部标题区 + 统计卡片行（4个）+ 快捷入口区 + 最近动态时间线。')

add_heading(doc, '4.2 页面元素', 2)
add_table(doc, ['区域', '元素', '说明'], [
    ['统计卡片', '进行中监测数', '管理员可见，数量=所有状态非PUBLISHED的监测数'],
    ['统计卡片', '本月待确认数', '管理员可见，当前CONFIRMING状态的监测数'],
    ['统计卡片', '已发布结果数', '管理员可见，PUBLISHED状态的监测总数'],
    ['统计卡片', '我的待办数', '收数人：待填写任务数；机构负责人：待确认数'],
    ['快捷入口', '发起月度监测', '管理员可见，跳转 /monitoring/create'],
    ['快捷入口', '填写绩效数据', '收数人可见，跳转 /collect'],
    ['快捷入口', '查看绩效报表', '所有人可见，跳转 /report'],
    ['快捷入口', '通知中心', '所有人可见，显示未读数 badge'],
    ['最近动态', '时间线列表', '显示最近10条操作记录，含时间、操作人、内容摘要'],
], col_widths=[2.5, 3, 6])

add_heading(doc, '4.3 关联API', 2)
add_table(doc, ['接口', '方法', '说明'], [
    ['GET /api/dashboard/stats', 'GET', '获取统计数据'],
    ['GET /api/dashboard/recent', 'GET', '获取最近动态'],
], col_widths=[4, 1.5, 6])

doc.add_page_break()
add_heading(doc, '5. 考核体系列表页 /system', 1)

add_page_header(doc, '5', '考核体系列表页', '/system', '管理员', '展示所有考核体系列表，支持搜索、分页、创建、编辑、删除操作。')

add_heading(doc, '5.1 布局', 2)
doc.add_paragraph('页面标题区（含"创建体系"按钮）+ 搜索筛选区 + 数据表格区（底部分页器）。')

add_heading(doc, '5.2 搜索筛选区', 2)
add_table(doc, ['元素', '类型', '说明'], [
    ['体系名称', 'Input.Search', '模糊搜索，支持按名称过滤'],
    ['状态筛选', 'Select', '全部 / 启用 / 禁用'],
    ['重置按钮', 'Button', '清空筛选条件'],
], col_widths=[3, 2.5, 6])

add_heading(doc, '5.3 数据表格列', 2)
add_table(doc, ['列名', '宽度', '说明'], [
    ['体系名称', '200px', '文本，超长时 tooltip 显示完整名称，点击跳转详情页'],
    ['描述', '300px', '文本，最多显示2行，超长省略'],
    ['机构数量', '100px', '数字，右对齐'],
    ['指标数量', '100px', '数字，右对齐'],
    ['状态', '100px', 'Tag：启用（绿色）/ 禁用（灰色）'],
    ['创建时间', '160px', '格式 YYYY-MM-DD HH:mm'],
    ['操作', '200px', '操作按钮组：编辑 / 查看 / 删除'],
], col_widths=[2.5, 1.5, 7])

add_heading(doc, '5.4 操作按钮', 2)
ops = [
    '创建体系：Primary 按钮，点击跳转 /system/create',
    '编辑：Text 按钮，点击跳转 /system/:id/edit，仅体系无进行中监测时可用',
    '查看：Text 按钮，点击跳转 /system/:id（只读详情）',
    '删除：Text 按钮（危险色），点击弹出确认框，删除后刷新列表',
]
for o in ops:
    doc.add_paragraph(o, style='List Bullet')

add_heading(doc, '5.5 空状态', 2)
doc.add_paragraph('表格无数据时，显示空状态插图 + 文字"暂无考核体系，立即创建"，下方显示"创建体系"按钮。')

add_heading(doc, '5.6 关联API', 2)
add_table(doc, ['接口', '方法', '说明'], [
    ['GET /api/systems', 'GET', '体系列表（分页+搜索）'],
    ['DELETE /api/systems/:id', 'DELETE', '删除体系'],
], col_widths=[4, 1.5, 6])

doc.add_page_break()
add_heading(doc, '6. 创建/编辑体系页 /system/create', 1)

add_page_header(doc, '6', '创建/编辑体系页', '/system/create 或 /system/:id/edit', '管理员', '上传Excel模版，系统解析并展示预览，确认无误后创建体系。编辑时模版不可修改。')

add_heading(doc, '6.1 布局', 2)
doc.add_paragraph('表单区：分为基础信息填写区 + Excel模版上传区 + 解析预览区 + 底部操作按钮区。编辑模式时隐藏模版上传区。')

add_heading(doc, '6.2 表单字段', 2)
add_table(doc, ['字段', '类型', '必填', '校验规则', '说明'], [
    ['体系名称', 'Input', '是', '最大100字符，非空', '体系名称'],
    ['体系描述', 'TextArea', '否', '最大500字符', '可选，描述该体系的适用范围'],
    ['是否需要审批', 'Switch', '否', '默认关闭', '开启后收数人提交需管理员审批才能入库'],
    ['Excel模版文件', 'Upload.Dragger', '是（创建时）', '后缀.xlsx，大小<=5MB', '拖拽上传区域，上传后显示文件名和大小'],
], col_widths=[2.5, 2.5, 1.5, 3.5, 3])

add_heading(doc, '6.3 解析预览区', 2)
doc.add_paragraph('Excel上传成功后，系统解析并展示四个Sheet的预览（只读表格，各最多显示前10行）：')
previews = [
    '模版页预览：列维度、类别、一级指标、二级指标、权重等17列',
    '机构页预览：机构名称、机构ID、分组名称、机构负责人',
    '数据收集页预览：收数指标名称、收数人、指标单位（展示前3行表头）',
    '参数页预览：参数名称、参数值',
]
for p in previews:
    doc.add_paragraph(p, style='List Bullet')

add_heading(doc, '6.4 错误处理', 2)
add_table(doc, ['错误类型', '提示方式', '处理方式'], [
    ['Sheet数量不足4个', 'Alert 错误提示', '阻止提交，提示缺失的Sheet名称'],
    ['模版页列不全或顺序不对', 'Alert 错误提示', '阻止提交，列出缺失的列'],
    ['机构页数据为空', 'Alert 错误提示', '阻止提交'],
    ['文件格式非Excel', 'Upload组件内置提示', '不允许选择该文件'],
    ['文件超过5MB', 'Upload组件内置提示', '不允许选择该文件'],
], col_widths=[3.5, 3.5, 4])

add_heading(doc, '6.5 操作按钮', 2)
add_table(doc, ['按钮', '类型', '行为'], [
    ['取消', 'Default Button', '返回列表页，有修改时弹窗确认'],
    ['上一步', 'Default Button', '编辑模式无此按钮'],
    ['创建体系', 'Primary Button', '校验表单+解析结果，提交创建请求'],
], col_widths=[2.5, 2.5, 6])

add_heading(doc, '6.6 关联API', 2)
add_table(doc, ['接口', '方法', '说明'], [
    ['POST /api/systems', 'POST', '创建体系（含文件上传）'],
    ['PUT /api/systems/:id', 'PUT', '编辑体系基本信息'],
    ['GET /api/systems/:id', 'GET', '获取体系详情用于编辑回填'],
], col_widths=[4, 1.5, 6])

doc.add_page_break()
add_heading(doc, '7. 月度监测列表页 /monitoring', 1)

add_page_header(doc, '7', '月度监测列表页', '/monitoring', '管理员', '展示所有月度监测记录，按状态筛选，支持发起新监测。')

add_heading(doc, '7.1 布局', 2)
doc.add_paragraph('页面标题区（含"发起监测"按钮）+ 搜索筛选区 + 数据表格区 + 底部分页器。')

add_heading(doc, '7.2 搜索筛选区', 2)
add_table(doc, ['元素', '类型', '说明'], [
    ['体系选择', 'Select', '下拉选择已有体系，支持搜索'],
    ['年份选择', 'Select', '下拉选择年份，范围：当前年份-2 ~ 当前年份+1'],
    ['月份选择', 'Select', '下拉选择月份 1-12'],
    ['状态筛选', 'Select', '全部 / 待发起 / 收数中 / 已截止 / 数据落库中 / 待确认 / 已发布'],
    ['重置', 'Button', '清空筛选条件'],
], col_widths=[3, 2.5, 5.5])

add_heading(doc, '7.3 数据表格列', 2)
add_table(doc, ['列名', '宽度', '说明'], [
    ['体系名称', '180px', '文本'],
    ['监测月份', '120px', '格式 YYYY年MM月'],
    ['状态', '130px', 'Tag，颜色按状态色彩规范'],
    ['收数截止时间', '170px', '格式 YYYY-MM-DD HH:mm'],
    ['进度', '150px', 'PROCESSING状态显示 Progress 进度条；其他状态显示"-"'],
    ['发起人', '100px', '文本'],
    ['发起时间', '160px', '格式 YYYY-MM-DD HH:mm'],
    ['操作', '180px', '查看详情 / 更多操作下拉菜单'],
], col_widths=[2.5, 1.5, 7])

add_heading(doc, '7.4 更多操作菜单', 2)
add_table(doc, ['操作', '可见条件', '行为'], [
    ['整体上传数据', 'COLLECTING状态', '弹出上传对话框'],
    ['手动截止', 'COLLECTING状态', '点击弹窗确认，确认后状态变为CLOSED'],
    ['查看确认进度', 'CONFIRMING状态', '显示确认进度弹窗'],
    ['发布结果', 'CONFIRMING状态 + 所有机构已确认', 'Primary按钮，点击弹窗确认后发布'],
    ['下载报告', 'PUBLISHED状态', '触发文件下载'],
], col_widths=[3, 3, 5])

add_heading(doc, '7.5 发起监测弹窗', 2)
add_table(doc, ['字段', '类型', '说明'], [
    ['选择体系', 'Select', '必选，下拉选择'],
    ['年份', 'InputNumber', '必选，默认当前年份，范围当前年份-2~+1'],
    ['月份', 'InputNumber', '必选，默认当前月份，范围1-12'],
    ['收数截止时间', 'DatePicker', '必选，支持时分选择，不早于当前时间+1小时'],
    ['是否需要审批', 'Switch', '默认继承体系的need_approval设置'],
], col_widths=[3, 2.5, 6])

add_heading(doc, '7.6 关联API', 2)
add_table(doc, ['接口', '方法', '说明'], [
    ['GET /api/monitorings', 'GET', '监测列表（分页+筛选）'],
    ['POST /api/monitorings', 'POST', '发起月度监测'],
    ['POST /api/monitorings/:id/close', 'POST', '手动截止收数'],
    ['POST /api/monitorings/:id/publish', 'POST', '管理员发布'],
], col_widths=[4, 1.5, 6])

doc.add_page_break()
add_heading(doc, '8. 监测详情页 /monitoring/:id', 1)

add_page_header(doc, '8', '监测详情页', '/monitoring/:id', '管理员', '展示特定月度监测的完整信息，包括数据收集情况、落库进度、确认进度、关联文件等。')

add_heading(doc, '8.1 布局', 2)
doc.add_paragraph('顶部步骤条（监测流程状态）+ 基本信息卡片 + Tab切换区（数据收集 / 落库进度 / 确认情况 / 关联文件）。')

add_heading(doc, '8.2 顶部步骤条', 2)
add_table(doc, ['步骤', '状态条件', '当前步骤高亮'], [
    ['1. 发起监测', 'PENDING', 'COLLECTING时高亮step1'],
    ['2. 数据收集', 'COLLECTING', 'CLOSED时高亮step2'],
    ['3. 数据落库', 'CLOSED / PROCESSING', 'CONFIRMING时高亮step3'],
    ['4. 负责人确认', 'CONFIRMING', 'PUBLISHED时高亮step4'],
    ['5. 发布结果', 'PUBLISHED', '已发布时高亮step5'],
], col_widths=[3, 4, 4])

add_heading(doc, '8.3 基本信息卡片', 2)
add_table(doc, ['字段', '值来源'], [
    ['体系名称', '来自关联体系'],
    ['监测月份', '来自月度监测记录'],
    ['当前状态', '来自月度监测记录'],
    ['收数截止时间', '来自月度监测记录'],
    ['进行中收数任务', '实时统计'],
    ['已完成收数任务', '实时统计'],
    ['数据落库进度', '仅PROCESSING状态显示Progress'],
], col_widths=[3, 8])

add_heading(doc, '8.4 Tab1 - 数据收集', 2)
doc.add_paragraph('COLLECTING状态时显示"整体上传"按钮 + 任务分发列表。')
add_table(doc, ['子Tab', '内容'], [
    ['收数任务列表', '按机构分组展示，每个机构展开显示其下所有指标的收数任务，含状态Tag'],
    ['整体上传记录', '展示所有管理员整体上传的历史记录，含上传时间和操作人'],
], col_widths=[3, 8])

add_table(doc, ['收数任务列表列', '说明'], [
    ['机构名称', '机构页中的机构名称'],
    ['指标名称', '二级指标名称'],
    ['收数人', '姓名/工号格式'],
    ['实际值', '已填写时显示数值，未填写显示"-"'],
    ['状态', 'Tag: 待填写(灰) / 已提交(蓝) / 已审批(绿) / 已驳回(红)'],
    ['提交时间', 'YYYY-MM-DD HH:mm'],
    ['操作', '查看 / 编辑（收数人自己的任务）'],
], col_widths=[3, 8])

add_heading(doc, '8.5 Tab2 - 数据落库进度（PROCESSING状态）', 2)
add_table(doc, ['元素', '说明'], [
    ['Progress 进度条', '实时显示 process_percent，按百分比推进'],
    ['处理状态标签', 'idle / processing / done / failed'],
    ['处理信息', '显示当前正在处理的内容，如"正在生成第3/10个机构的报告"'],
    ['错误信息', 'failed时显示 process_msg（错误详情）'],
    ['自动刷新', '页面每5秒自动调用接口刷新进度'],
], col_widths=[3, 8])

add_heading(doc, '8.6 Tab3 - 确认情况（CONFIRMING状态）', 2)
add_table(doc, ['确认情况列', '说明'], [
    ['机构名称', '机构名称'],
    ['机构负责人', '姓名/工号'],
    ['确认状态', 'Tag: 待确认(橙) / 已确认(绿)'],
    ['确认时间', 'YYYY-MM-DD HH:mm，未确认显示"-"'],
    ['确认备注', '负责人填写的备注，未确认显示"-"'],
    ['操作', '查看本机构数据（点击跳转该机构报表）'],
], col_widths=[3, 8])

add_heading(doc, '8.7 Tab4 - 关联文件', 2)
doc.add_paragraph('以文件树形式展示MinIO中该监测的所有生成文件：')
add_table(doc, ['目录层级', '内容'], [
    ['{体系ID}/{年份}/{月份}/', '该监测下所有机构月度报告文件'],
    ['每个机构文件夹', 'report_{年月}.xlsx（可下载）'],
    ['发布前', '文件可下载，验证用户权限'],
    ['发布后', '文件公开下载'],
], col_widths=[3, 8])

add_heading(doc, '8.8 关联API', 2)
add_table(doc, ['接口', '方法', '说明'], [
    ['GET /api/monitorings/:id', 'GET', '获取监测详情'],
    ['GET /api/monitorings/:id/process/status', 'GET', '获取数据落库进度'],
    ['GET /api/monitorings/:id/collect/tasks', 'GET', '获取收数任务列表'],
    ['POST /api/monitorings/:id/collect/upload', 'POST', '整体上传数据收集页Excel'],
    ['GET /api/monitorings/:id/files', 'GET', '获取关联文件列表'],
    ['GET /api/reports/download/:monitoringId/:institutionId', 'GET', '下载特定机构的月度报告'],
], col_widths=[4, 1.5, 6])

doc.add_page_break()
add_heading(doc, '9. 数据填写页（收数人） /collect', 1)

add_page_header(doc, '9', '数据填写页（收数人）', '/collect', '收数人', '收数人查看自己被分配的指标任务，填写实际值后一次性提交。')

add_heading(doc, '9.1 布局', 2)
doc.add_paragraph('页面标题（含当前用户信息）+ 任务筛选区（按体系、月份）+ 任务列表 + 底部提交按钮。')

add_heading(doc, '9.2 任务筛选区', 2)
add_table(doc, ['元素', '类型', '说明'], [
    ['体系选择', 'Select', '仅显示当前用户有收数任务的体系'],
    ['月份选择', 'Select', '可选月份，仅显示有进行中任务的月份'],
    ['状态筛选', 'Select', '全部 / 待填写 / 已提交'],
], col_widths=[3, 2.5, 6])

add_heading(doc, '9.3 任务卡片列表', 2)
doc.add_paragraph('每个任务以卡片形式展示，卡片内包含：')
add_table(doc, ['卡片内容', '说明'], [
    ['任务标题', '"{体系名} - {月份} - {指标名}"'],
    ['所属机构', '机构名称'],
    ['指标单位', '单位'],
    ['全年目标', '数值'],
    ['进度目标', '数值（月度进度目标）'],
    ['实际值输入框', 'InputNumber，精度根据单位自动适配，placeholder="请输入实际完成值"'],
    ['状态标识', 'Tag：待填写 / 已提交'],
    ['提交时间', '已提交时显示'],
], col_widths=[3, 8])

add_heading(doc, '9.4 交互逻辑', 2)
steps = [
    '收数人进入页面后，自动查询当前用户所有进行中的收数任务',
    '同一机构的所有指标可以一起填写，只需一次提交（批量提交）',
    '按机构分组展示，每个机构组内含所有该机构下该收数人负责的指标',
    '填写实际值后，点击"提交"按钮，弹出确认框（显示填写的数据摘要）',
    '确认后调用 PUT /api/monitorings/:id/collect/tasks/:taskId 提交数据',
    '提交成功后，任务状态变为"已提交"，输入框变为只读',
    '如需修改（管理员驳回后），需联系管理员重置状态',
    '下载按钮：每个任务卡片提供"下载Excel"按钮，下载仅含数据收集页的Excel模板',
    '上传按钮：支持上传已填写好的Excel文件，覆盖当前填写的数据（需二次确认）',
]
for s in steps:
    doc.add_paragraph(s, style='List Bullet')

add_heading(doc, '9.5 空状态', 2)
doc.add_paragraph('无收数任务时显示："当前没有需要填写的绩效数据" + 插图。')

add_heading(doc, '9.6 关联API', 2)
add_table(doc, ['接口', '方法', '说明'], [
    ['GET /api/collect/tasks', 'GET', '获取当前收数人的所有任务（按机构分组）'],
    ['PUT /api/monitorings/:id/collect/tasks/:taskId', 'PUT', '提交单个任务数据'],
    ['POST /api/monitorings/:id/collect/upload', 'POST', '上传Excel批量提交'],
    ['GET /api/collect/template/download', 'GET', '下载仅含数据收集页的Excel'],
], col_widths=[4, 1.5, 6])

doc.add_page_break()
add_heading(doc, '10. 绩效报表页 /report', 1)

add_page_header(doc, '10', '绩效报表页', '/report', '管理员、机构负责人（发布后）', '展示特定机构在特定月份的绩效数据，支持筛选、Excel导出和图表可视化。机构负责人在发布前只能看自己机构。')

add_heading(doc, '10.1 布局', 2)
doc.add_paragraph('顶部筛选器栏 + Tab切换（数据表格 / 可视化图表）+ 底部操作栏（导出Excel、打印）。')

add_heading(doc, '10.2 筛选器栏', 2)
add_table(doc, ['元素', '类型', '说明'], [
    ['考核体系', 'Select', '必选'],
    ['年份', 'Select', '必选，默认当前年份'],
    ['月份', 'Select', '必选，默认当前月份'],
    ['机构（管理员可选，负责人隐藏）', 'Select（支持多选）', '可选，不选则默认展示所有机构'],
    ['分组筛选', 'Select', '可选，选择后仅展示该分组内的机构'],
    ['重置按钮', 'Button', '恢复默认筛选条件'],
], col_widths=[3.5, 2.5, 5])

add_heading(doc, '10.3 Tab1 - 数据表格', 2)
add_table(doc, ['列名', '说明'], [
    ['机构名称', '机构名称，分组列可折叠'],
    ['分组', '分组名称'],
    ['维度', '指标维度'],
    ['类别', '指标类别'],
    ['一级指标', '一级指标名称'],
    ['二级指标', '二级指标名称'],
    ['实际值', '数值，右对齐'],
    ['全年完成率', '百分比，保留2位小数'],
    ['进度完成率', '百分比，保留2位小数'],
    ['指标百分制得分', '数值，保留4位小数'],
    ['指标权重得分', '数值，保留4位小数'],
    ['维度得分', '数值，保留4位小数'],
    ['总得分', '数值，保留4位小数，加粗显示'],
], col_widths=[3, 8])

add_heading(doc, '10.4 Tab2 - 可视化图表', 2)
add_table(doc, ['图表类型', '说明', '交互'], [
    ['分组机构排名柱状图', '横向柱状图，展示所有机构总得分排名', '鼠标悬停显示详情，点击钻取到机构详情'],
    ['维度得分雷达图', '多维度得分对比（每个机构一条线）', '可切换只看单一机构'],
    ['指标完成率进度条', '各指标实际值 vs 进度目标', '支持按机构筛选'],
    ['类别得分饼图', '各类别权重得分占比', '点击扇区钻取到类别明细'],
], col_widths=[3, 5, 3])

add_heading(doc, '10.5 机构负责人权限控制', 2)
doc.add_paragraph('机构负责人在PUBLISHED状态前，筛选器中机构字段默认选中且置灰为其关联机构，不可切换。PUBLISHED状态后可查看所有机构数据。')

add_heading(doc, '10.6 关联API', 2)
add_table(doc, ['接口', '方法', '说明'], [
    ['GET /api/reports/institution/:id', 'GET', '获取特定机构报表数据'],
    ['GET /api/reports/summary', 'GET', '获取所有机构汇总数据'],
    ['GET /api/reports/export/excel', 'GET', '导出Excel报表（根据筛选条件）'],
], col_widths=[4, 1.5, 6])

doc.add_page_break()
add_heading(doc, '11. 总览对比页 /report/overview', 1)

add_page_header(doc, '11', '总览对比页', '/report/overview', '管理员、机构负责人（发布后）', '综合展示某体系某月份下所有机构的绩效对比情况，包括分组排名、雷达图、关键指标对比图等。')

add_heading(doc, '11.1 布局', 2)
doc.add_paragraph('顶部筛选器栏（体系+年月） + 主内容区（2列布局：左侧图表区 + 右侧排名列表区）。')

add_heading(doc, '11.2 筛选器栏', 2)
add_table(doc, ['元素', '类型', '说明'], [
    ['考核体系', 'Select', '必选'],
    ['年份', 'Select', '必选'],
    ['月份', 'Select', '必选'],
    ['分组筛选', 'Select', '可选，用于切换不同的分组对比视图'],
], col_widths=[3, 2.5, 5.5])

add_heading(doc, '11.3 左侧图表区（60%宽度）', 2)
add_table(doc, ['图表', '类型', '说明'], [
    ['总得分排名', '横向柱状图', '所有机构按总得分降序排列，颜色按分组着色'],
    ['维度得分雷达图', '雷达图', '多机构叠加显示，每机构一条线，支持图例切换显示/隐藏'],
    ['关键指标对比', '分组柱状图', '选择3-5个核心指标，横轴为机构，纵轴为数值的分组柱状图'],
    ['分组得分汇总', '堆叠柱状图', '展示各分组平均得分对比'],
], col_widths=[3, 2.5, 5.5])

add_heading(doc, '11.4 右侧排名列表区（40%宽度）', 2)
add_table(doc, ['内容', '说明'], [
    ['分组Tab', '每个分组一个Tab，无分组时显示"全部机构"'],
    ['排名列表', '该分社内所有机构按总得分排名，含机构名、分组、得分、名次'],
    ['Top3高亮', '前3名高亮显示（金银铜色）'],
    ['我的机构', '当前用户的关联机构用特殊颜色标识（机构负责人）'],
    ['点击跳转', '点击某机构行，跳转到该机构详情报表（/report）并筛选到该机构'],
], col_widths=[3, 8])

add_heading(doc, '11.5 关联API', 2)
add_table(doc, ['接口', '方法', '说明'], [
    ['GET /api/reports/summary', 'GET', '获取总览数据（含各机构得分、排名、分组信息）'],
    ['GET /api/reports/ranking', 'GET', '获取分组排名数据'],
], col_widths=[4, 1.5, 6])

doc.add_page_break()
add_heading(doc, '12. 通知中心页 /notifications', 1)

add_page_header(doc, '12', '通知中心页', '/notifications', '所有登录用户', '展示当前用户收到的所有站内通知，支持标记已读。')

add_heading(doc, '12.1 布局', 2)
doc.add_paragraph('Tab切换区（全部 / 未读）+ 通知列表 + 底部分页器。Header区显示未读总数Badge。')

add_heading(doc, '12.2 Tab区', 2)
add_table(doc, ['Tab', '说明'], [
    ['全部', '显示所有通知'],
    ['未读', '仅显示未读通知'],
], col_widths=[2.5, 8.5])

add_heading(doc, '12.3 通知卡片', 2)
add_table(doc, ['元素', '说明'], [
    ['通知图标', '根据类型显示：收数任务(蓝色)、截止提醒(橙色)、确认提醒(紫色)、发布通知(绿色)'],
    ['通知标题', '加粗文本，如"【收数任务】您有新的数据待填写"'],
    ['通知内容', '正文内容，如"体系：XXX，月份：2026年03月，请于XXX前完成"'],
    ['时间', '相对时间，如"3分钟前"、"2小时前"、"3天前"，鼠标悬停显示精确时间'],
    ['状态标识', '未读时左侧有蓝色竖条标识，已读无'],
    ['操作', '未读时显示"标为已读"，已读时无操作'],
], col_widths=[3, 8])

add_heading(doc, '12.4 交互逻辑', 2)
steps = [
    '页面加载时自动调用接口获取通知列表（未读优先排序）',
    '点击通知卡片跳转到对应页面（如：收数任务通知 -> /collect）',
    '点击"标为已读"，调用接口标记单条通知为已读，左侧蓝色竖条消失',
    '"全部标为已读"按钮（未读Tab下显示），一键已读所有通知',
    'Header区通知图标显示未读数Badge，点击直接跳转本页面',
    '有新通知时，通过前端轮询（每60秒）检测并更新Badge数字',
]
for s in steps:
    doc.add_paragraph(s, style='List Bullet')

add_heading(doc, '12.5 关联API', 2)
add_table(doc, ['接口', '方法', '说明'], [
    ['GET /api/notifications', 'GET', '获取通知列表（分页）'],
    ['PUT /api/notifications/:id/read', 'PUT', '标记单条已读'],
    ['PUT /api/notifications/read-all', 'PUT', '标记全部已读'],
    ['GET /api/notifications/unread-count', 'GET', '获取未读数量（Header Badge用）'],
], col_widths=[4, 1.5, 6])

doc.add_page_break()
add_heading(doc, '13. 通用组件规范', 1)

add_heading(doc, '13.1 确认对话框', 2)
doc.add_paragraph('所有危险操作（删除、发布、截止）需弹出 Modal.confirm 确认框，包含：操作描述 + 取消/确认按钮。发布和截止操作在确认框中显示操作摘要（如"即将发布2026年3月的绩效数据，发布后数据将公开"）。')

add_heading(doc, '13.2 成功/失败提示', 2)
add_table(doc, ['场景', '组件', '行为'], [
    ['操作成功', 'message.success', '自动消失，时长3秒'],
    ['操作失败', 'message.error', '自动消失，时长5秒'],
    ['表单校验失败', 'Form item 错误提示', '红色文字显示在字段下方'],
    ['网络错误', 'message.error', '显示"网络错误，请稍后重试"'],
], col_widths=[3, 3, 5])

add_heading(doc, '13.3 加载状态', 2)
add_table(doc, ['场景', '处理方式'], [
    ['页面初始加载', '页面内容区显示 Spin 加载动画，遮罩整个内容区'],
    ['表格数据加载', '表格区域内显示 Spin，覆盖表格但不遮挡表头'],
    ['按钮操作中', 'Button 设置 loading=true，显示 spinner 文字变为"处理中..."'],
    ['文件上传中', 'Upload 组件内置 Progress 进度条'],
], col_widths=[3, 8])

add_heading(doc, '13.4 权限控制', 2)
doc.add_paragraph('前端通过路由守卫（React Router Auth Guard）控制页面级访问权限：')
auth_rules = [
    '未登录访问任何页面 -> 跳转 /login',
    '收数人访问 /system 或 /monitoring -> 跳转 /dashboard 并提示"无权限"',
    '机构负责人访问 /system 或 /monitoring -> 跳转 /dashboard 并提示"无权限"',
    '机构负责人访问 /report 时如监测未PUBLISHED -> 仅展示其关联机构数据',
]
for r in auth_rules:
    doc.add_paragraph(r, style='List Bullet')

add_heading(doc, '13.5 分页组件', 2)
doc.add_paragraph('统一使用 Ant Design Pagination 组件，默认配置：current=1, pageSize=10, showSizeChanger, showTotal。')

doc.add_page_break()
add_heading(doc, '14. 响应式与无障碍', 1)

add_heading(doc, '14.1 响应式策略', 2)
add_table(doc, ['断点', '布局变化'], [
    ['>= 1200px', '标准布局，Sider 240px展开'],
    ['768px - 1199px', 'Sider 64px折叠，仅显示图标'],
    ['< 768px', 'Sider 收起为 Drawer 抽屉，移动端布局'],
], col_widths=[3, 8])

add_heading(doc, '14.2 无障碍规范', 2)
items = [
    '所有表单输入必须关联 label（htmlFor）',
    '所有按钮和交互元素必须有可访问的名称',
    '图片必须包含 alt 属性描述',
    '颜色对比度满足 WCAG 2.1 AA 标准',
    '键盘可完全操作（Tab导航、回车确认、Esc取消）',
    '焦点状态（focus）样式清晰可见',
    '页面标题（<title>）需包含页面名称',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()
add_heading(doc, '附录：各页面与API对照表', 1)
add_table(doc, ['页面', '关联API'], [
    ['登录页 /login', 'POST /api/auth/login'],
    ['首页 /dashboard', 'GET /api/dashboard/stats, GET /api/dashboard/recent'],
    ['考核体系列表 /system', 'GET /api/systems, DELETE /api/systems/:id'],
    ['创建体系 /system/create', 'POST /api/systems'],
    ['编辑体系 /system/:id/edit', 'GET /api/systems/:id, PUT /api/systems/:id'],
    ['月度监测列表 /monitoring', 'GET /api/monitorings, POST /api/monitorings, POST /api/monitorings/:id/close, POST /api/monitorings/:id/publish'],
    ['监测详情 /monitoring/:id', 'GET /api/monitorings/:id, GET /api/monitorings/:id/process/status, GET /api/monitorings/:id/collect/tasks, POST /api/monitorings/:id/collect/upload, GET /api/monitorings/:id/files'],
    ['数据填写 /collect', 'GET /api/collect/tasks, PUT /api/monitorings/:id/collect/tasks/:taskId, POST /api/monitorings/:id/collect/upload'],
    ['绩效报表 /report', 'GET /api/reports/institution/:id, GET /api/reports/summary, GET /api/reports/export/excel'],
    ['总览对比 /report/overview', 'GET /api/reports/summary, GET /api/reports/ranking'],
    ['通知中心 /notifications', 'GET /api/notifications, PUT /api/notifications/:id/read, PUT /api/notifications/read-all, GET /api/notifications/unread-count'],
], col_widths=[4, 7])

doc.save('/Users/zhaoyu/Desktop/机构绩效管理系统_页面设计文档.docx')
print('Page design doc generated successfully')
